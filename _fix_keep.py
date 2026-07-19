# -*- coding: utf-8 -*-
"""Keep each table on a single page + minor title-page degrees (match reference).
   All our tables are <=9 rows, so keepNext can never overflow a page.
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_keep(p_el, lines=True):
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    tags = ("w:keepNext", "w:keepLines") if lines else ("w:keepNext",)
    for tag in tags:
        if pPr.find(qn(tag)) is None:
            # keepNext/keepLines should come early in pPr; append is fine for Word
            pPr.append(OxmlElement(tag))


def ptext(e):
    return "".join(t.text or "" for t in e.findall(".//" + qn("w:t"))).strip()


def set_para_text(p, newtext):
    runs = p.runs
    if runs:
        runs[0].text = newtext
        for r in runs[1:]:
            r.text = ""


def process(fn):
    d = Document(fn)
    body = d.element.body
    ch = list(body.iterchildren())
    kept_tables = 0
    kept_caps = 0
    for i, e in enumerate(ch):
        if e.tag == qn("w:tbl"):
            # keepNext+keepLines on every paragraph inside the table
            for p_el in e.findall(".//" + qn("w:p")):
                set_keep(p_el, lines=True)
            kept_tables += 1
            # caption directly above -> keep with table
            j = i - 1
            while j >= 0:
                if ch[j].tag == qn("w:p") and ptext(ch[j]):
                    if ptext(ch[j]).startswith("Таблица"):
                        set_keep(ch[j], lines=True)
                        kept_caps += 1
                    break
                j -= 1

    # title-page academic degrees (reference lists both as к.т.н.)
    deg = 0
    for p in d.paragraphs[:35]:
        t = p.text.strip()
        if t == "Вагнер Виктория Игоревна":
            set_para_text(p, "к.т.н. Вагнер Виктория Игоревна")
            deg += 1
        elif t.startswith("Нормоконтролер") and "к.т.н" not in t:
            set_para_text(p, "Нормоконтролер: к.т.н. Чалова Екатерина Игорьевна")
            deg += 1

    # count heading entries that feed the TOC (level 1 + level 2)
    h1 = sum(1 for p in d.paragraphs if p.style.name == "Heading 1")
    h2 = sum(1 for p in d.paragraphs if p.style.name == "Heading 2")

    d.save(fn)
    return {
        "kept_tables": kept_tables,
        "kept_caps": kept_caps,
        "degrees": deg,
        "toc_entries": h1 + h2,
        "h1": h1,
        "h2": h2,
    }


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "->", process(fn))
