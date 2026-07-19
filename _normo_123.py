# -*- coding: utf-8 -*-
"""Normocontrol items 1,2,3 on the working copy.

1. "Выводы по N главе": remove the numbered heading line, prepend
   "Выводы по N главе." as a bold lead-in to the first conclusion paragraph.
   Also remove these lines from the manual СОДЕРЖАНИЕ.
2. Remove empty paragraphs sitting strictly between two body-text paragraphs.
3. Report tables whose number is never referenced in the body text.
"""
import io
import re
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_нормоконтроль.docx"
doc = Document(SRC)
report = io.StringIO()

paras = doc.paragraphs

# locate intro / содержание boundaries
intro_idx = None
toc_idx = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t == "СОДЕРЖАНИЕ" and toc_idx is None:
        toc_idx = i
    if t == "ВВЕДЕНИЕ" and intro_idx is None:
        intro_idx = i


def del_paragraph(p):
    p._element.getparent().remove(p._element)


def prepend_bold(target_p, text):
    new_r = target_p.add_run(text)
    new_r.bold = True
    r_el = new_r._element
    target_p._p.remove(r_el)
    ppr = target_p._p.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(r_el)
    else:
        target_p._p.insert(0, r_el)


# ---------- 1. Выводы по главе ----------
vyvod_heading = re.compile(r"^\s*\d+\.\d+\s+(Выводы по .+? главе)\s*$")
removed_vyvod = 0
# collect target paragraph objects first (indices shift on delete, so work via elements)
to_process = []
for i, p in enumerate(paras):
    if intro_idx is not None and i <= intro_idx:
        continue  # skip TOC region for the body change
    m = vyvod_heading.match(p.text.strip())
    if m:
        # find next non-empty paragraph
        nxt = None
        for j in range(i + 1, len(paras)):
            if paras[j].text.strip():
                nxt = paras[j]
                break
        to_process.append((p, nxt, m.group(1)))

for heading_p, nxt_p, phrase in to_process:
    if nxt_p is not None:
        prepend_bold(nxt_p, phrase + ". ")
    del_paragraph(heading_p)
    removed_vyvod += 1

# remove выводы lines from TOC (before intro)
removed_toc = 0
for p in list(doc.paragraphs):
    # re-evaluate index region by checking we're before ВВЕДЕНИЕ text appears
    pass
# simpler: iterate, stop at ВВЕДЕНИЕ
for p in list(doc.paragraphs):
    t = p.text.strip()
    if t == "ВВЕДЕНИЕ":
        break
    if re.search(r"Выводы по .+? главе", t) and "...." in t:
        del_paragraph(p)
        removed_toc += 1

# ---------- 2. Empty paragraphs between text ----------
paras = doc.paragraphs  # refresh
removed_empty = 0
caption_re = re.compile(r"^(таблица|рисунок)\b", re.IGNORECASE)
for i in range(1, len(paras) - 1):
    p = paras[i]
    if p.text.strip():
        continue
    prev = paras[i - 1].text.strip()
    nxt = paras[i + 1].text.strip()
    if not prev or not nxt:
        continue
    # both neighbors normal text (not heading number, not caption)
    if re.match(r"^\d", prev) or nxt[:1].isdigit():
        continue
    if caption_re.match(prev) or caption_re.match(nxt):
        continue
    # this empty paragraph: ensure it's truly empty (no images/breaks)
    if p._element.findall(".//" + qn("w:drawing")) or p._element.findall(
        ".//" + qn("w:br")
    ):
        continue
    del_paragraph(p)
    removed_empty += 1

# ---------- 3. Tables without in-text reference ----------
full_text = "\n".join(p.text for p in doc.paragraphs)
# table captions: "Таблица N — ..."
cap_re = re.compile(r"^Таблица\s+(\d+)\s*[—\-–]\s*(.+)$")
table_caps = []
for p in doc.paragraphs:
    m = cap_re.match(p.text.strip())
    if m:
        table_caps.append((int(m.group(1)), m.group(2)[:60]))

report.write("=== ТАБЛИЦЫ БЕЗ ССЫЛКИ В ТЕКСТЕ ===\n")
report.write("(надо дописать в тексте до таблицы: «...приведены в таблице N»)\n\n")
unref = []
for num, name in table_caps:
    # reference patterns: таблиц* N (but not the caption line "Таблица N —")
    pat = re.compile(r"табл[а-яё]*\.?\s*" + str(num) + r"\b", re.IGNORECASE)
    # count matches NOT part of a caption
    refs = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if cap_re.match(t):
            continue  # skip caption itself
        if pat.search(t):
            refs += 1
    if refs == 0:
        unref.append((num, name))

if unref:
    for num, name in unref:
        report.write(f"  Таблица {num} — {name}  →  ССЫЛКИ НЕТ\n")
else:
    report.write("  Все таблицы упомянуты в тексте.\n")
report.write(f"\nИтого таблиц: {len(table_caps)}, без ссылки: {len(unref)}\n")

doc.save(SRC)

with open(
    r"C:/Users/egor3/Desktop/smart-price/_tables_no_ref.txt", "w", encoding="utf-8"
) as f:
    f.write(report.getvalue())

print(f"1) Выводы переоформлены: {removed_vyvod}, удалено из оглавления: {removed_toc}")
print(f"2) Пустых абзацев удалено: {removed_empty}")
print(f"3) Таблиц без ссылки: {len(unref)} (см. _tables_no_ref.txt)")
print("Saved.")
