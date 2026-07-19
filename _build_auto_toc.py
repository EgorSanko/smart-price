# -*- coding: utf-8 -*-
"""(B) Build a real Word auto-TOC.

1. Redefine Heading 1/2/3 styles to GOST (Times New Roman 14, bold, black;
   H1 centered, H2/H3 left).
2. Apply Heading 1 to chapter-level headings (1-4 ALL CAPS, ВВЕДЕНИЕ,
   ЗАКЛЮЧЕНИЕ, СПИСОК…, ПРИЛОЖЕНИЕ); Heading 2 to x.x sections.
   Keeps existing direct formatting (page breaks, indent, spacing).
3. Delete the manual СОДЕРЖАНИЕ entries; insert a TOC field after the
   СОДЕРЖАНИЕ title. Word fills page numbers on "Update field" (F9).
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_нормоконтроль.docx"
doc = Document(SRC)


# ---------- 1. Redefine heading styles to GOST ----------
def style_gost(style, *, center):
    f = style.font
    f.name = "Times New Roman"
    f.size = Pt(14)
    f.bold = True
    f.color.rgb = RGBColor(0, 0, 0)
    # also set East Asian / cs font so it sticks
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), "Times New Roman")
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True


for name, center in [("Heading 1", True), ("Heading 2", False), ("Heading 3", False)]:
    try:
        style_gost(doc.styles[name], center=center)
    except KeyError:
        pass

# ---------- locate boundaries ----------
intro_idx = soder_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "СОДЕРЖАНИЕ" and soder_idx is None:
        soder_idx = i
    if t == "ВВЕДЕНИЕ" and intro_idx is None:
        intro_idx = i

# ---------- 2. Apply heading styles in body (from ВВЕДЕНИЕ onward) ----------
h1_special = {"ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"}
sec_re = re.compile(r"^\d+\.\d+\s+\S")
chap_re = re.compile(r"^[1-4]\s+[А-ЯЁ]")

h1 = h2 = 0
for i, p in enumerate(doc.paragraphs):
    if intro_idx is None or i < intro_idx:
        continue
    t = p.text.strip()
    if not t:
        continue
    is_h1 = (
        (t in h1_special)
        or t.startswith("ПРИЛОЖЕНИЕ")
        or (chap_re.match(t) and t.isupper() and len(t) < 80)
    )
    is_h2 = bool(sec_re.match(t)) and len(t) < 90
    if is_h1:
        p.style = doc.styles["Heading 1"]
        h1 += 1
    elif is_h2:
        p.style = doc.styles["Heading 2"]
        h2 += 1

# ---------- 3. Delete manual TOC entries, insert TOC field ----------
# delete paragraphs strictly between СОДЕРЖАНИЕ and ВВЕДЕНИЕ
soder_p = doc.paragraphs[soder_idx]
removed = 0
for p in list(doc.paragraphs):
    # only those after СОДЕРЖАНИЕ and before ВВЕДЕНИЕ
    pass
# rebuild by element walk
collecting = False
for p in list(doc.paragraphs):
    t = p.text.strip()
    if p._p is soder_p._p:
        collecting = True
        continue
    if collecting:
        if t == "ВВЕДЕНИЕ":
            break
        # delete this TOC entry paragraph
        p._p.getparent().remove(p._p)
        removed += 1


# Build TOC field paragraph and insert right after СОДЕРЖАНИЕ
def make_toc_paragraph():
    p = OxmlElement("w:p")
    # field begin
    r1 = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r1.append(fc1)
    p.append(r1)
    # instr
    r2 = OxmlElement("w:r")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = ' TOC \\o "1-2" \\h \\z \\u '
    r2.append(it)
    p.append(r2)
    # separate
    r3 = OxmlElement("w:r")
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r3.append(fc2)
    p.append(r3)
    # placeholder text
    r4 = OxmlElement("w:r")
    t4 = OxmlElement("w:t")
    t4.set(qn("xml:space"), "preserve")
    t4.text = "Обновите оглавление в Word: ПКМ по этому полю → «Обновить поле» → «Обновить целиком»."
    r4.append(t4)
    p.append(r4)
    # end
    r5 = OxmlElement("w:r")
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r5.append(fc3)
    p.append(r5)
    return p


toc_p = make_toc_paragraph()
soder_p._p.addnext(toc_p)

doc.save(SRC)
print(f"Heading 1 applied: {h1}, Heading 2 applied: {h2}")
print(f"Manual TOC entries removed: {removed}")
print("TOC field inserted after СОДЕРЖАНИЕ.")
