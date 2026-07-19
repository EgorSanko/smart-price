# -*- coding: utf-8 -*-
"""Script 8 — remove in-text table-reference announcements (нормоконтроль).
   Delete pure announcement paragraphs; trim ref lead-ins/sentences in mixed ones.
   Does NOT touch DB-table names ('в таблице users' etc. — no numeric index).
"""
import re
from docx import Document
from docx.oxml.ns import qn

ANN = re.compile(
    r"(привед[её]н[аоы]?|показан[аоы]?|представлен[аоы]?|изображ[её]н[аоы]?|"
    r"отображ[её]н[аоы]?|продемонстрирован[аоы]?)\s+в\s+таблице\s+\d",
    re.I,
)
ENDS_TBL = re.compile(r"в\s+таблице\s+\d+\s*\.$")

MIXED = [
    ("Как видно из таблицы 1, исследуемые", "Исследуемые"),
    ("Представленная в таблице 2 матрица", "Матрица"),
    ("Результаты представлены в таблице 11. ", ""),
]


def para_replace(p, old, new):
    runs = p.runs
    texts = [r.text for r in runs]
    full = "".join(texts)
    idx = full.find(old)
    if idx < 0:
        return False
    end = idx + len(old)
    pos = 0
    for r, txt in zip(runs, texts):
        rs, re_ = pos, pos + len(txt)
        pos = re_
        if re_ <= idx or rs >= end:
            continue
        a = max(idx, rs) - rs
        b = min(end, re_) - rs
        ins = new if (rs <= idx < re_) else ""
        r.text = txt[:a] + ins + txt[b:]
    return True


def process(fn):
    d = Document(fn)
    # 1. delete pure announcement paragraphs
    to_delete = []
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        sents = [s for s in re.split(r"(?<=[.!?])\s+", t) if s.strip()]
        if len(sents) == 1 and ANN.search(t) and ENDS_TBL.search(t):
            to_delete.append(p)
    for p in to_delete:
        p._element.getparent().remove(p._element)
    # 2. mixed: trim the reference part
    mixed_done = 0
    for old, new in MIXED:
        for p in d.paragraphs:
            if old in p.text:
                if para_replace(p, old, new):
                    mixed_done += 1
                break
    d.save(fn)
    return {"pure_deleted": len(to_delete), "mixed_trimmed": mixed_done}


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "->", process(fn))
