# -*- coding: utf-8 -*-
"""Create a "clean" docx for antiplagiat check.

What we strip (per instructor request):
  - title page (everything before the РЕФЕРАТ heading)
  - table of contents (СОДЕРЖАНИЕ block)
  - references (СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ to end of file)
  - appendices (ПРИЛОЖЕНИЕ ... to end)
  - all images
  - all tables (data + caption lines like "Таблица 11 — …")
  - all code listings (Листинг N — ... and the code block that follows)
  - all figure captions ("Рисунок 1.1 — …")

What we keep:
  - РЕФЕРАТ
  - ВВЕДЕНИЕ
  - main body chapters 1–4
  - ЗАКЛЮЧЕНИЕ

Strategy: work directly on the body XML, removing <w:p>/<w:tbl> elements
that match exclusion rules. Then walk the surviving paragraphs and drop
any inline drawings.
"""
import re
import shutil
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
DST = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_антиплагиат.docx"

shutil.copy(SRC, DST)
doc = Document(DST)
body = doc.element.body

# All children of body in document order: paragraphs (w:p), tables (w:tbl),
# section properties (w:sectPr), etc.
children = list(body.iterchildren())


# Locate boundary paragraphs by exact-ish text match
def para_text(elem):
    if elem.tag != qn("w:p"):
        return None
    texts = elem.findall(".//" + qn("w:t"))
    return "".join(t.text or "" for t in texts).strip()


# Build a list of (index, kind, text) for traversal
boundaries = {
    "referat_start": None,  # keep from here
    "toc_start": None,  # remove from here
    "toc_end": None,  # keep from here (ВВЕДЕНИЕ)
    "references_start": None,  # remove from here to end (or until appendix is already in remove zone)
    "appendix_start": None,  # remove from here to end
}

# NOTE on the dot-leader trick:
#   Inside the TOC each entry looks like "ВВЕДЕНИЕ ......... 7", so the
#   pure-text match on "ВВЕДЕНИЕ" hits both the TOC line and the real
#   chapter heading. To pick only the heading, we additionally require
#   the absence of ".....". Same for СПИСОК… and ПРИЛОЖЕНИЕ.


def is_heading_not_toc_entry(text, prefix):
    if not text.startswith(prefix):
        return False
    # TOC entries always carry a dot leader. Real headings never do.
    return "...." not in text


for i, c in enumerate(children):
    t = para_text(c) or ""
    if t.startswith("РЕФЕРАТ") and boundaries["referat_start"] is None:
        boundaries["referat_start"] = i
    elif t.startswith("СОДЕРЖАНИЕ") and boundaries["toc_start"] is None:
        boundaries["toc_start"] = i
    elif (
        is_heading_not_toc_entry(t, "ВВЕДЕНИЕ")
        and boundaries["toc_end"] is None
        and boundaries["toc_start"] is not None
    ):
        boundaries["toc_end"] = i
    elif (
        is_heading_not_toc_entry(t, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        and boundaries["references_start"] is None
    ):
        boundaries["references_start"] = i
    elif (
        is_heading_not_toc_entry(t, "ПРИЛОЖЕНИЕ")
        and boundaries["appendix_start"] is None
    ):
        boundaries["appendix_start"] = i

print(f"Boundaries: {boundaries}")

# Build set of indices to remove
to_remove = set()

# 1) Everything before РЕФЕРАТ (title page block)
if boundaries["referat_start"] is not None:
    for i in range(0, boundaries["referat_start"]):
        to_remove.add(i)

# 2) TOC block — from СОДЕРЖАНИЕ to (exclusive) ВВЕДЕНИЕ
if boundaries["toc_start"] is not None:
    end_idx = (
        boundaries["toc_end"]
        if boundaries["toc_end"] is not None
        else boundaries["toc_start"] + 60
    )
    for i in range(boundaries["toc_start"], end_idx):
        to_remove.add(i)

# 3) References + Appendices — from earliest of references/appendix start to end
tail_starts = [
    b
    for b in (boundaries["references_start"], boundaries["appendix_start"])
    if b is not None
]
if tail_starts:
    tail_start = min(tail_starts)
    for i in range(tail_start, len(children)):
        to_remove.add(i)

# 4) Per-element exclusion rules — applied to surviving elements only
TABLE_CAPTION_RE = re.compile(r"^\s*Таблица\s+\d", re.IGNORECASE)
FIGURE_CAPTION_RE = re.compile(r"^\s*Рисунок\s+\d", re.IGNORECASE)
LISTING_CAPTION_RE = re.compile(r"^\s*Листинг\s+\d", re.IGNORECASE)


# Code-block heuristic: paragraph whose font is monospace OR
# the line contains many non-cyrillic structural chars (braces, semicolons,
# operators) and few words. We mark such paragraphs for removal.
def looks_like_code(text):
    if not text or len(text) < 4:
        return False
    # Common code tokens: { } ; => -> def class function import
    code_tokens = (
        "def ",
        "class ",
        "import ",
        "return ",
        "async ",
        "=>",
        "->",
        "==",
        "!=",
        "&&",
        "||",
        "function(",
        "function ",
    )
    if any(tok in text for tok in code_tokens):
        return True
    # Lots of braces or semicolons relative to length
    structural = sum(1 for ch in text if ch in "{}[];=:|>")
    if structural >= 4 and structural / max(1, len(text)) > 0.05:
        return True
    return False


# Mark caption and listing paragraphs (and a few following lines for listings)
removed_listing_block = False
listing_block_lines_left = 0

for i, c in enumerate(children):
    if i in to_remove:
        continue
    if c.tag == qn("w:tbl"):
        to_remove.add(i)
        continue
    if c.tag != qn("w:p"):
        continue
    t = para_text(c) or ""
    if not t:
        continue
    if TABLE_CAPTION_RE.match(t) or FIGURE_CAPTION_RE.match(t):
        to_remove.add(i)
        continue
    if LISTING_CAPTION_RE.match(t):
        to_remove.add(i)
        # mark the next 30 paragraphs to scan for code
        listing_block_lines_left = 30
        continue
    if listing_block_lines_left > 0:
        # While inside a listing zone, drop monospace/code-like paragraphs.
        # Stop on first paragraph that looks like normal prose (>= 50 chars
        # with mostly Cyrillic letters and no code tokens).
        cyr_letters = sum(
            1 for ch in t if "а" <= ch.lower() <= "я" or ch.lower() == "ё"
        )
        if looks_like_code(t) or len(t) < 100 and cyr_letters < len(t) * 0.5:
            to_remove.add(i)
            listing_block_lines_left -= 1
        else:
            listing_block_lines_left = 0  # back to prose

# Apply removals
for i in sorted(to_remove, reverse=True):
    body.remove(children[i])

# 5) Strip inline drawings (images) from surviving paragraphs
DRAWING_TAGS = (qn("w:drawing"), qn("w:pict"), qn("w:object"))
removed_drawings = 0
for p in doc.paragraphs:
    for r in list(p.runs):
        for tag in DRAWING_TAGS:
            for d in r._element.findall(".//" + tag):
                d.getparent().remove(d)
                removed_drawings += 1

doc.save(DST)
print(f"Removed: {len(to_remove)} elements, {removed_drawings} drawings")
print(f"Saved to {DST}")

# Diagnostics: how many paragraphs survived
doc2 = Document(DST)
print(f"Surviving paragraphs: {len(doc2.paragraphs)}, tables: {len(doc2.tables)}")
