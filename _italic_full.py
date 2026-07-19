# -*- coding: utf-8 -*-
"""Comprehensive italic/mono pass over the WHOLE document.

Covers everything the first pass skipped: РЕФЕРАТ (before ВВЕДЕНИЕ),
headings, captions, lead-ins, table cells. Foreign (Latin) terms -> italic;
code identifiers -> Consolas (no italic). Idempotent (re-asserts italic
val=true so previously-broken w:i val=0 runs get fixed too).

Skips ONLY the auto-TOC region (paragraphs between СОДЕРЖАНИЕ and ВВЕДЕНИЕ)
so Word's generated TOC field is left for Word to regenerate.
"""
import io
import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_нормоконтроль.docx"
doc = Document(SRC)

TOKEN = re.compile(r"[A-Za-z_][\w]*(?:[.\-/][\w]+)*(?:\([^)]*\))?(?:\[[^\]]*\])?")
CURATED_CODE = {
    "asyncio.gather",
    "EventSourceResponse",
    "EAliceOfferCard",
    "EAliceOffer",
    "EAliceOfferList",
    "app.ml",
    "app.celery_app",
    "app.main",
    "sse_starlette",
}


def classify(tok):
    if len(tok) < 2:
        return None
    if "_" in tok or "(" in tok or "[" in tok or tok in CURATED_CODE:
        return "code"
    if not re.search(r"[A-Za-z]", tok):
        return None
    return "term"


def segment(text):
    segs, pos = [], 0
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            segs.append((text[pos : m.start()], None))
        segs.append((m.group(0), classify(m.group(0))))
        pos = m.end()
    if pos < len(text):
        segs.append((text[pos:], None))
    return segs


def run_is_simple(r):
    for bad in (
        "w:br",
        "w:drawing",
        "w:pict",
        "w:tab",
        "w:fldChar",
        "w:object",
        "w:instrText",
    ):
        if r.findall(".//" + qn(bad)):
            return False
    return len(r.findall(qn("w:t"))) <= 1


def _rpr(r):
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    return rPr


def set_italic(r):
    rPr = _rpr(r)
    i = rPr.find(qn("w:i"))
    if i is None:
        i = rPr.makeelement(qn("w:i"), {})
        rPr.append(i)
    i.set(qn("w:val"), "true")
    ics = rPr.find(qn("w:iCs"))
    if ics is None:
        ics = rPr.makeelement(qn("w:iCs"), {})
        rPr.append(ics)
    ics.set(qn("w:val"), "true")


def set_mono(r):
    rPr = _rpr(r)
    for tag in ("w:i", "w:iCs"):
        e = rPr.find(qn(tag))
        if e is not None:
            rPr.remove(e)
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), "Consolas")


def is_italic(r):
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        return False
    i = rPr.find(qn("w:i"))
    return i is not None and i.get(qn("w:val")) not in ("0", "false")


def is_mono(r):
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        return False
    rf = rPr.find(qn("w:rFonts"))
    return rf is not None and rf.get(qn("w:ascii")) == "Consolas"


stats = {"italic": 0, "mono": 0}


def process(p):
    for r in list(p.runs):
        re_el = r._element
        if not run_is_simple(re_el):
            continue
        text = r.text
        if not text or not TOKEN.search(text):
            continue
        segs = segment(text)
        if not any(k for _, k in segs):
            continue
        # already fully styled single-token run?
        if len(segs) == 1:
            k = segs[0][1]
            if k == "term":
                if not is_italic(re_el):
                    set_italic(re_el)
                    stats["italic"] += 1
            elif k == "code":
                if not is_mono(re_el):
                    set_mono(re_el)
                    stats["mono"] += 1
            continue
        parent = re_el.getparent()
        idx = list(parent).index(re_el)
        for off, (st, k) in enumerate(segs):
            nr = deepcopy(re_el)
            t = nr.find(qn("w:t"))
            if t is None:
                t = nr.makeelement(qn("w:t"), {})
                nr.append(t)
            t.text = st
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            if k == "term":
                set_italic(nr)
                stats["italic"] += 1
            elif k == "code":
                set_mono(nr)
                stats["mono"] += 1
            parent.insert(idx + off, nr)
        parent.remove(re_el)


# boundaries: skip TOC region (between СОДЕРЖАНИЕ and ВВЕДЕНИЕ)
soder_idx = intro_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "СОДЕРЖАНИЕ" and soder_idx is None:
        soder_idx = i
    if t == "ВВЕДЕНИЕ" and intro_idx is None:
        intro_idx = i

for i, p in enumerate(doc.paragraphs):
    # skip title page (paras before РЕФЕРАТ have no Latin worth; but safe to process — skip TOC only)
    if soder_idx is not None and intro_idx is not None and soder_idx < i < intro_idx:
        continue  # TOC field region
    process(p)

# tables too
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process(p)

doc.save(SRC)
print(f'Italic applied: {stats["italic"]}, mono applied: {stats["mono"]}')
