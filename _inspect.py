# -*- coding: utf-8 -*-
"""Inspect the diploma — write findings to UTF-8 file (Windows console can't handle Cyrillic)."""
from docx import Document
import io

doc = Document(r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx")

out = io.StringIO()

# 1. Find section 3.13 area with paragraph index, style, bold flags per run
out.write("=== AREA 1: paragraphs around index 415 (section 3.13) ===\n")
for i in range(410, 460):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    runs_info = [(r.bold, r.text[:40]) for r in p.runs]
    out.write(f"[{i}] style={p.style.name!r} text={p.text[:100]!r}\n")
    for j, (b, t) in enumerate(runs_info):
        out.write(f"    run{j}: bold={b} text={t!r}\n")
    out.write("\n")

out.write('\n=== AREA 2: search "союзн" everywhere ===\n')
for i, p in enumerate(doc.paragraphs):
    if "союзн" in p.text.lower():
        out.write(f"[{i}] {p.text[:200]!r}\n")

out.write('\n=== AREA 3: tables for "союзн" ===\n')
for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            if "союзн" in cell.text.lower():
                out.write(f"table{ti} row{ri} cell{ci}: {cell.text[:200]!r}\n")

out.write("\n=== AREA 4: TOC area (paragraphs 25-100) ===\n")
for i in range(20, 100):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    out.write(f"[{i}] style={p.style.name!r} text={p.text[:120]!r}\n")

with open(
    r"C:/Users/egor3/Desktop/smart-price/_inspect.txt", "w", encoding="utf-8"
) as f:
    f.write(out.getvalue())

print("OK, written to _inspect.txt")
