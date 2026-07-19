# -*- coding: utf-8 -*-
"""Targeted search for word-level typos in the diploma."""
import re
import io
from docx import Document

doc = Document(r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx")


def iter_all_paragraphs(d):
    for i, p in enumerate(d.paragraphs):
        yield ("P", i, p.text)
    for ti, t in enumerate(d.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield (f"t{ti}r{ri}c{ci}p{pi}", None, p.text)


# Suspicious patterns
PATTERNS = [
    # Россииипиа, Россииипиа, etc — repeated "ии" or "иа" inside Russian words
    (r"\b\w*Росси[а-я]{2,}\b", "Russia-glue"),
    (r"\b\w*Беларус[а-я]{2,}\b", "Belarus-glue"),
    (r"\b\w*россии\w+", "россии+"),
    (r"\b\w*беларуси\w+", "беларуси+"),
    # Cyrillic word longer than 25 chars
    (r"\b[А-ЯЁа-яё-]{25,}\b", "long-cyr-word"),
    # double letter glue: пиа, иап, etc inside short bigrams of suffixes
    (r"(России|Беларуси)[а-яё]", "suffix-glue"),
]

# Russian suffix endings allowed for Russia/Belarus
RUSSIA_VALID = {
    "России",
    "Россия",
    "Российская",
    "Российской",
    "Российскую",
    "Россию",
    "Российских",
    "Российский",
    "Россиян",
    "Российскую",
    "Российские",
    "Россиянами",
    "Россиянка",
    "Россиянам",
}
BELARUS_VALID = {
    "Беларусь",
    "Беларуси",
    "Беларусью",
    "Беларусам",
    "Белорусский",
    "Белорусская",
    "Белорусских",
    "Беларусь.",
    "Беларусь,",
    "Беларуси.",
    "Беларуси,",
    "Беларуси;",
    "Беларуси:",
    "Беларусь:",
    "Беларусь;",
}

out = io.StringIO()
seen = set()

for kind, idx, text in iter_all_paragraphs(doc):
    if not text or not text.strip():
        continue
    # check long words
    for m in re.finditer(r"[А-ЯЁа-яё-]{20,}", text):
        word = m.group(0)
        if word in seen:
            continue
        seen.add(word)
        # OK if matches valid forms
        if word in RUSSIA_VALID or word in BELARUS_VALID:
            continue
        ctx_start = max(0, m.start() - 30)
        ctx_end = min(len(text), m.end() + 30)
        out.write(f"[{kind}{idx}] LONG-WORD: {word!r}\n")
        out.write(f"  ctx: ...{text[ctx_start:ctx_end]}...\n\n")

    # check Russia/Belarus glue
    for m in re.finditer(r"(России|Беларуси|Россия|Беларусь)([А-ЯЁа-яё]+)", text):
        full = m.group(0)
        if full in seen:
            continue
        seen.add(full)
        ctx_start = max(0, m.start() - 30)
        ctx_end = min(len(text), m.end() + 30)
        out.write(f"[{kind}{idx}] GLUE: {full!r}\n")
        out.write(f"  ctx: ...{text[ctx_start:ctx_end]}...\n\n")

with open(r"C:/Users/egor3/Desktop/smart-price/_typos.txt", "w", encoding="utf-8") as f:
    f.write(out.getvalue())
print(f"Found {len(seen)} unique suspicious tokens.")
