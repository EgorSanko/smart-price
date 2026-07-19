# -*- coding: utf-8 -*-
"""Inspect hyperlinks, body bold, tables, listings."""
import io
import re
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
doc = Document(SRC)
out = io.StringIO()

# 1. Hyperlinks in document
out.write("=== HYPERLINKS (w:hyperlink elements) ===\n")
body = doc.element.body
hyperlinks = body.findall(".//" + qn("w:hyperlink"))
out.write(f"Total w:hyperlink elements: {len(hyperlinks)}\n")

# 2. References section — find СПИСОК... and dump a few paragraphs
out.write("\n=== REFERENCES SAMPLE (after СПИСОК ИСПОЛЬЗОВАННЫХ) ===\n")
ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("СПИСОК ИСПОЛЬЗОВАННЫХ"):
        ref_start = i
        break
if ref_start is not None:
    for p in doc.paragraphs[ref_start + 1 : ref_start + 8]:
        if p.text.strip():
            runs_info = [
                (
                    r.text[:30],
                    r.font.underline,
                    str(r.font.color.rgb)
                    if r.font.color and r.font.color.rgb
                    else None,
                )
                for r in p.runs
            ]
            out.write(f"  PARA: {p.text[:90]!r}\n")
            for txt, ul, color in runs_info:
                out.write(f"      run underline={ul} color={color} text={txt!r}\n")

# 3. Count bold runs in body (non-heading) paragraphs
heading_re = re.compile(r"^\s*\d+(\.\d+){0,2}\s+\S")
body_bold_count = 0
body_bold_samples = []
for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    is_heading = heading_re.match(t) or (t.isupper() and len(t) < 60)
    if is_heading:
        continue
    for r in p.runs:
        if r.bold and r.text.strip():
            body_bold_count += 1
            if len(body_bold_samples) < 25:
                body_bold_samples.append(r.text.strip()[:50])
            break
out.write(
    f"\n=== BODY (non-heading) paragraphs with bold runs: {body_bold_count} ===\n"
)
for s in body_bold_samples:
    out.write(f"  {s!r}\n")

# 4. Tables
out.write(f"\n=== TABLES: {len(doc.tables)} ===\n")
for ti, t in enumerate(doc.tables[:5]):
    out.write(f"  Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols\n")

# 5. Listings / code: find "Листинг" captions
out.write("\n=== LISTING CAPTIONS ===\n")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.lower().startswith("листинг"):
        out.write(f"  [{i}] {t[:80]!r}\n")

# 6. Figure / Table captions
out.write("\n=== FIGURE/TABLE CAPTIONS (count) ===\n")
fig = sum(1 for p in doc.paragraphs if p.text.strip().lower().startswith("рисунок"))
tab = sum(1 for p in doc.paragraphs if p.text.strip().lower().startswith("таблица"))
out.write(f"  Рисунок-captions: {fig}\n  Таблица-captions: {tab}\n")

with open(
    r"C:/Users/egor3/Desktop/smart-price/_format_inspect2.txt", "w", encoding="utf-8"
) as f:
    f.write(out.getvalue())
print("OK inspect2")
