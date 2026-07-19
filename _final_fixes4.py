# -*- coding: utf-8 -*-
"""
Final fixes v4 — addresses defense reviewer's feedback.

Three changes:
1. Section 3.13.1–3.13.5 BODY paragraphs are erroneously bold; restore normal weight
   while keeping subsection HEADERS (3.13, 3.13.1, ...) bold.
2. "Союзном государстве России и Беларуси" → "едином торговом пространстве
   России и Беларуси" (par. 81). Reviewer rejected the political union framing.
3. Replace placeholder "—" in the manual TOC (par. 31–75) with real page numbers
   extracted from rendered ДИПЛОМ.pdf via _extract_pages.py.
"""
import re
from docx import Document

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
doc = Document(SRC)

# -----------------------------------------------------------------------------
# Real page numbers (extracted from rendered PDF)
# -----------------------------------------------------------------------------
# Map TOC line prefix (text before the dots) -> page number
TOC_PAGES = {
    "ВВЕДЕНИЕ": 7,
    "1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ И СУЩЕСТВУЮЩИХ РЕШЕНИЙ": 14,
    "1.1 Рынок электронной коммерции и маркетплейсов": 14,
    "1.2 Анализ существующих решений": 20,
    "1.3 Технологический обзор": 24,
    "1.4 Выводы по первой главе": 30,
    "2 ПРОЕКТИРОВАНИЕ СИСТЕМЫ МЕТАПОИСКА": 32,
    "2.1 Анализ требований": 32,
    "2.2 Архитектурное проектирование": 34,
    "2.3 Проектирование базы данных": 38,
    "2.4 Проектирование API": 40,
    "2.5 Проектирование пользовательского интерфейса": 42,
    "2.6 Проектирование развёртывания": 46,
    "2.7 Проектирование обработки ошибок и отказоустойчивости": 47,
    "2.8 Проектирование кэширования": 48,
    "2.9 Выводы по второй главе": 48,
    "3 РЕАЛИЗАЦИЯ СИСТЕМЫ МЕТАПОИСКА": 50,
    "3.1 Реализация парсеров маркетплейсов": 50,
    "3.2 Реализация AI-коррекции запросов": 52,
    "3.3 Реализация гибридной системы фильтрации": 53,
    "3.4 Реализация AI-ассистента": 56,
    "3.5 Реализация SSE-стриминга": 58,
    "3.6 Реализация фронтенда": 60,
    "3.7 Интеграция с Gemini через OpenRouter": 62,
    "3.8 Реализация системы аутентификации": 63,
    "3.9 Реализация прокси-сервера изображений": 65,
    "3.10 Реализация модуля истории цен": 65,
    "3.11 Реализация AI-анализа цен": 66,
    "3.12 Развёртывание и DevOps": 69,
    "3.13 Реализация функции «Найти дешевле»": 70,
    "3.14 Выводы по третьей главе": 72,
    "4 ТЕСТИРОВАНИЕ И АНАЛИЗ РЕЗУЛЬТАТОВ": 74,
    "4.1 Методология тестирования": 74,
    "4.2 Результаты функционального тестирования": 75,
    "4.3 Результаты тестирования AI-ассистента": 77,
    "4.4 Результаты тестирования производительности": 78,
    "4.5 Сквозное тестирование интерфейса": 80,
    "4.6 Тестирование адаптивной вёрстки": 85,
    "4.7 Сравнение с существующими решениями": 86,
    "4.8 Экономический анализ эксплуатации": 87,
    "4.9 Выводы по четвёртой главе": 89,
    "ЗАКЛЮЧЕНИЕ": 91,
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ": 94,
    "ПРИЛОЖЕНИЕ А — Структура каталогов проекта": 96,
    "ПРИЛОЖЕНИЕ Б — Примеры ответов API": 101,
}

stats = {"bold_unset": 0, "union_replaced": 0, "toc_filled": 0, "toc_unmatched": []}

# -----------------------------------------------------------------------------
# Pass 1: union-state replacement (par. 81)
# -----------------------------------------------------------------------------
UNION_PATTERNS = [
    (
        "Союзном государстве России и Беларуси",
        "едином торговом пространстве России и Беларуси",
    ),
    (
        "Союзного государства России и Беларуси",
        "единого торгового пространства России и Беларуси",
    ),
    (
        "Союзное государство России и Беларуси",
        "единое торговое пространство России и Беларуси",
    ),
    (
        "союзном государстве России и Беларуси",
        "едином торговом пространстве России и Беларуси",
    ),
    ("Союзном государстве", "едином торговом пространстве"),
    ("Союзного государства", "единого торгового пространства"),
]


def apply_run_aware(paragraph, old, new):
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


for p in doc.paragraphs:
    for old, new in UNION_PATTERNS:
        if apply_run_aware(p, old, new):
            stats["union_replaced"] += 1
            break  # don't double-process

# -----------------------------------------------------------------------------
# Pass 2: un-bold body paragraphs of 3.13.1 — 3.13.5
# Heading paragraphs (start with "3.13", "3.13.1" etc.) stay bold; body becomes
# normal weight. Strategy: walk paragraphs sequentially in 3.13 range, detect
# headings vs body by regex on the text.
# -----------------------------------------------------------------------------
HEADING_PATTERN = re.compile(r"^\s*\d+(\.\d+){0,2}\s")


def is_heading_text(text):
    t = text.strip()
    if HEADING_PATTERN.match(t):
        return True
    # Also treat top-level chapter words as headings
    if t in ("ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        return True
    if t.startswith("ПРИЛОЖЕНИЕ "):
        return True
    return False


# Locate the 3.13 area: from "3.13 Реализация функции поиска более выгодных…"
# until we hit "3.14" or "4 ТЕСТИРОВАНИЕ" — these stop the un-bolding.
in_section_313 = False
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith("3.13 Реализация функции поиска") or text.startswith(
        "3.13 Реализация функции «"
    ):
        # heading 3.13 itself — keep bold, mark section start
        in_section_313 = True
        continue
    if in_section_313:
        if text.startswith("3.14") or text.startswith("4 ") or text.startswith("4\t"):
            in_section_313 = False
            continue
        if is_heading_text(text):
            # 3.13.1, 3.13.2, ... — keep bold (it's a subheading)
            continue
        # body paragraph inside 3.13 → un-bold all runs
        for r in p.runs:
            if r.bold:
                r.bold = False
                stats["bold_unset"] += 1

# -----------------------------------------------------------------------------
# Pass 3: TOC page numbers — replace trailing "—" with real page number,
# preserving leading dots. Format target: "TEXT ....... PAGE"
# -----------------------------------------------------------------------------
LINE_TOTAL = 70  # target total line width, similar to original


def make_toc_line(text, page):
    base = text.rstrip()
    page_str = str(page)
    # Compute dots so total length stays around LINE_TOTAL
    # Format: "{text} {dots} {page}" — separated by single spaces around dots.
    available = LINE_TOTAL - len(base) - 1 - 1 - len(page_str)
    dots = "." * max(3, available)
    return f"{base} {dots} {page_str}"


# Build a normalised lookup: strip leading whitespace, strip trailing "..." and "—"
def normalise_toc_text(txt):
    t = txt.strip()
    # remove trailing dot-leader and dash placeholder
    t = re.sub(r"\s*\.{3,}\s*[—–-]?\s*$", "", t)
    t = re.sub(r"\s*[—–-]\s*$", "", t)
    return t.strip()


for i, p in enumerate(doc.paragraphs):
    if i < 25 or i > 85:
        continue
    raw = p.text
    if not raw.strip():
        continue
    # only process lines that contain dot-leader pattern or end with "—"
    if (
        "...." not in raw
        and not raw.rstrip().endswith("—")
        and not raw.rstrip().endswith("-")
    ):
        continue
    norm = normalise_toc_text(raw)
    page = TOC_PAGES.get(norm)
    if page is None:
        # Try a tolerant lookup by prefix
        for key, pg in TOC_PAGES.items():
            if norm.startswith(key[:30]) or key.startswith(norm[:30]):
                page = pg
                break
    if page is None:
        stats["toc_unmatched"].append(norm)
        continue
    new_text = make_toc_line(norm, page)
    # Apply (collapse runs)
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    stats["toc_filled"] += 1

# -----------------------------------------------------------------------------
doc.save(SRC)
print("=== STATS ===")
for k, v in stats.items():
    print(f"{k}: {v}")
print("Saved.")
