# -*- coding: utf-8 -*-
"""Script 9 — figure spacing + footer 14pt + title underline hugging text."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

NAMES = ("Санько", "Вагнер", "Чалова")


def has_drawing(p):
    return p._element.find(".//" + qn("w:drawing")) is not None


def ptext(p):
    return p.text.strip()


def process(fn):
    d = Document(fn)

    # 1. figures: text -> gap -> image -> caption -> gap -> text
    body = d.element.body
    children = [c for c in body.iterchildren() if c.tag == qn("w:p")]
    figs = 0
    for i, el in enumerate(children):
        p = Paragraph(el, d)
        if not has_drawing(p):
            continue
        pf = p.paragraph_format
        pf.space_before = Pt(12)  # gap above image (between text and figure)
        pf.space_after = Pt(0)
        pf.keep_with_next = True  # keep image with its caption
        # caption = next paragraph starting with 'Рисунок'
        if i + 1 < len(children):
            cap = Paragraph(children[i + 1], d)
            if ptext(cap).startswith("Рисунок"):
                cpf = cap.paragraph_format
                cpf.space_before = Pt(0)
                cpf.space_after = Pt(12)  # gap below caption
        figs += 1

    # 2. footer page numbers -> 14 pt, Times New Roman
    for sec in d.sections:
        for p in sec.footer.paragraphs:
            for r in p.runs:
                r.font.size = Pt(14)
                r.font.name = "Times New Roman"
                rpr = r._element.get_or_add_rPr()
                rf = rpr.find(qn("w:rFonts"))
                if rf is None:
                    from docx.oxml import OxmlElement

                    rf = OxmlElement("w:rFonts")
                    rpr.insert(0, rf)
                for a in ("ascii", "hAnsi", "cs", "eastAsia"):
                    rf.set(qn("w:" + a), "Times New Roman")

    # 3. title signature lines: line right under text (single spacing, no after-gap)
    sig = 0
    for p in d.paragraphs[:30]:
        if any(n in p.text for n in NAMES):
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.line_spacing = 1.0
            pf.space_after = Pt(0)
            pf.space_before = Pt(8)
            # tighten the border distance
            pPr = p._element.find(qn("w:pPr"))
            pBdr = pPr.find(qn("w:pBdr")) if pPr is not None else None
            if pBdr is not None:
                bottom = pBdr.find(qn("w:bottom"))
                if bottom is not None:
                    bottom.set(qn("w:space"), "1")
            sig += 1

    d.save(fn)
    return {"figures": figs, "signatures": sig}


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "->", process(fn))
