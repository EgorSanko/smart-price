# -*- coding: utf-8 -*-
"""
Final fixes v3 — align diploma text with actual production configuration.
Reality check (measured 2026-04-19 from SSE stream /api/v1/live-search/stream):
  ACTIVE sources: onliner, yandex, regard, worlddevices, oneclick, biggeek
  NOT ACTIVE: wildberries (not in sources list at all)
Count: 6 parsers, not 5.
"""
import re
from docx import Document

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
doc = Document(SRC)

# Ordered replacements — more specific first to avoid cascading.
REPLACEMENTS = [
    # Parser count: 5 → 6
    ("пять парсеров", "шесть парсеров"),
    ("Пять парсеров", "Шесть парсеров"),
    ("пяти парсеров", "шести парсеров"),
    ("пять источников", "шесть источников"),
    ("Пять источников", "Шесть источников"),
    ("пяти источников", "шести источников"),
    ("5 парсеров", "6 парсеров"),
    ("5 источников", "6 источников"),
    ("5 активных", "6 активных"),
    ("5 отобраны", "6 отобраны"),
    ("ИСТОЧНИКОВ ПАРАЛЛЕЛЬНО\n5", "ИСТОЧНИКОВ ПАРАЛЛЕЛЬНО\n6"),
    # Full parser list: replace WB with 1click+BigGeek
    (
        "Onliner.by, Яндекс Маркет, Wildberries, Регард, World Devices",
        "Onliner.by, Яндекс Маркет, Регард, World Devices, 1click, BigGeek",
    ),
    (
        "Onliner, Я.Маркет, Wildberries, Регард, World Devices",
        "Onliner, Я.Маркет, Регард, World Devices, 1click, BigGeek",
    ),
    (
        "Onliner.by, Яндекс Маркет, Wildberries, Регард, World Devices.",
        "Onliner.by, Яндекс Маркет, Регард, World Devices, 1click, BigGeek.",
    ),
    # TTFR realistic numbers (measured from prod)
    ("TTFR < 2 с", "TTFR ≈ 5,8 с до первого батча (1,4 с до визуального отклика)"),
    ("TTFR < 5 с", "TTFR ≈ 5,8 с до первого батча (1,4 с до визуального отклика)"),
    # Container count: 5 → 6
    ("5 контейнеров", "6 контейнеров"),
]


def apply_run_aware(paragraph, old, new):
    """Replace across runs — collapse into first run, empty the rest."""
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


changed = 0
for p in doc.paragraphs:
    for old, new in REPLACEMENTS:
        if apply_run_aware(p, old, new):
            changed += 1
            print(
                f'  [OK] "{old[:40]}..." -> "{new[:40]}..."'.encode(
                    "ascii", "replace"
                ).decode()
            )

# Also walk tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for old, new in REPLACEMENTS:
                    if apply_run_aware(p, old, new):
                        changed += 1
                        print(
                            f'  [OK][TABLE] "{old[:40]}..." -> "{new[:40]}..."'.encode(
                                "ascii", "replace"
                            ).decode()
                        )

# Additionally: remove standalone "Wildberries" list items / mentions in parser contexts
# (leave Wildberries intact in roadmap / competitor analysis contexts)
WB_REMOVAL_PATTERNS = [
    r"·\s*Wildberries\s*87\s*offers[^\n]*",
    r"parser\.wildberries\s+rate-limit",
]
for p in doc.paragraphs:
    for pattern in WB_REMOVAL_PATTERNS:
        full = "".join(r.text for r in p.runs)
        new_full = re.sub(pattern, "", full)
        if new_full != full:
            if p.runs:
                p.runs[0].text = new_full
                for r in p.runs[1:]:
                    r.text = ""
            changed += 1
            print(f"  [OK] Removed WB-specific pattern: {pattern}")

doc.save(SRC)
print(f"\nTotal replacements: {changed}")
print("Saved.")
