# -*- coding: utf-8 -*-
"""Cleanup: remove AliExpress stub, rewrite "четыре из семи" line."""
import sys, re
from docx import Document

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
doc = Document(SRC)

# Remove AliExpress stub bullet
for p in list(doc.paragraphs):
    if "AliExpress" in p.text and "заглушка" in p.text:
        p._element.getparent().remove(p._element)
        print("Removed AliExpress stub")
        break

# Rewrite "четыре парсера из семи зарегистрированных ..." fully
for p in doc.paragraphs:
    if "активны четыре парсера из семи" in p.text or "четыре парсера из семи" in p.text:
        new_text = (
            "Все пять парсеров находятся в рабочей конфигурации: Onliner.by, "
            "Яндекс Маркет, Wildberries, Регард, World Devices. Парсеры запускаются "
            "параллельно при каждом поисковом запросе, что минимизирует суммарное "
            "время сбора данных."
        )
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
        print("Rewrote parser-count paragraph")
        break

doc.save(SRC)
print("Saved.")
