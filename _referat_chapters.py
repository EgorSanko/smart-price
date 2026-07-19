# -*- coding: utf-8 -*-
"""Restructure РЕФЕРАТ: trim verbose blocks, add 4 chapter summaries.
Keeps total length ~same so it still fits 2 pages. Then italicize Latin.
Applied to both docx files.
"""
import re
from copy import deepcopy
from docx.oxml.ns import qn
from docx import Document

CEL = (
    "Целью работы является проектирование и разработка веб-платформы метапоиска "
    "товаров на маркетплейсах Беларуси и России с AI-анализом цен на базе языковой "
    "модели Gemini 2.5 Flash, объединяющей шесть источников в единый интерфейс с "
    "потоковым отображением результатов через технологию Server-Sent Events."
)
MET = (
    "Методы исследования: системный анализ, объектно-ориентированное проектирование, "
    "веб-скрапинг с обходом антибот-защиты, обработка естественного языка на основе "
    "больших языковых моделей, экспериментальное тестирование и статистический анализ "
    "данных о ценах."
)
CH1 = (
    "В первой главе проанализированы рынок электронной коммерции России и Беларуси, "
    "существующие агрегаторы цен и технологический стек, обоснована актуальность "
    "разработки."
)
CH2 = (
    "Во второй главе спроектирована архитектура системы: модель базы данных, REST API "
    "и SSE-протокол потокового поиска, пользовательский интерфейс и схема развёртывания "
    "в Docker."
)
CH3 = (
    "В третьей главе реализованы шесть парсеров маркетплейсов, AI-коррекция запросов, "
    "гибридная фильтрация результатов, AI-ассистент и функция «Найти дешевле» на основе "
    "реверс-инжиниринга WebSocket-протокола Яндекс.Алисы."
)
CH4 = (
    "В четвёртой главе выполнено функциональное и нагрузочное тестирование системы, "
    "измерены метрики производительности и проведён экономический анализ эксплуатации."
)
RES = (
    "В результате разработана и развёрнута в промышленную эксплуатацию система Smart Price "
    "(smrt-price.ru): шесть параллельных парсеров, AI-ассистент и автоматический анализ "
    "цен; время до первого результата составило 1,4 секунды."
)
OBL = (
    "Область применения: электронная коммерция, системы поддержки принятия решений "
    "потребителей, агрегация и кроссрегиональный мониторинг цен на товары массового спроса."
)

# ----- italic helper -----
TOKEN = re.compile(r"[A-Za-z_][\w]*(?:[.\-/][\w]+)*")
TERMS = {
    "AI",
    "Gemini",
    "Server-Sent",
    "Events",
    "REST",
    "API",
    "SSE",
    "Docker",
    "Smart",
    "Price",
    "smrt-price.ru",
    "WebSocket",
    "FastAPI",
    "Next.js",
    "OpenRouter",
    "PostgreSQL",
    "Flash",
}


def set_ital(r):
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    for tag in ("w:i", "w:iCs"):
        e = rPr.find(qn(tag))
        if e is None:
            e = rPr.makeelement(qn(tag), {})
            rPr.append(e)
        e.set(qn("w:val"), "true")


def italize(p):
    for r in list(p.runs):
        re_el = r._element
        text = r.text
        if not text or not TOKEN.search(text):
            continue
        segs = []
        pos = 0
        for m in TOKEN.finditer(text):
            if m.start() > pos:
                segs.append((text[pos : m.start()], False))
            segs.append((m.group(0), m.group(0) in TERMS))
            pos = m.end()
        if pos < len(text):
            segs.append((text[pos:], False))
        if not any(it for _, it in segs):
            continue
        parent = re_el.getparent()
        idx = list(parent).index(re_el)
        for off, (st, it) in enumerate(segs):
            nr = deepcopy(re_el)
            t = nr.find(qn("w:t"))
            if t is None:
                t = nr.makeelement(qn("w:t"), {})
                nr.append(t)
            t.text = st
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            if it:
                set_ital(nr)
            parent.insert(idx + off, nr)
        parent.remove(re_el)


def settext(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def process(fn):
    d = Document(fn)
    # locate реферат body paragraphs
    cel = met = res = obl = None
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith("Целью") and cel is None:
            cel = p
        elif t.startswith("Методы исследования") and met is None:
            met = p
        elif (
            t.startswith("Основные результаты") or t.startswith("В результате")
        ) and res is None:
            res = p
        elif t.startswith("Область применения") and obl is None:
            obl = p
    settext(cel, CEL)
    settext(met, MET)
    settext(res, RES)
    settext(obl, OBL)
    # insert 4 chapter summaries before результаты
    for txt in [CH1, CH2, CH3, CH4]:
        res.insert_paragraph_before(txt)
    # italicize Latin in all реферат body paragraphs (from РЕФЕРАТ to ОГЛАВЛЕНИЕ)
    rf = sod = None
    for i, p in enumerate(d.paragraphs):
        if p.text.strip() == "РЕФЕРАТ":
            rf = i
        if p.text.strip() in ("ОГЛАВЛЕНИЕ", "СОДЕРЖАНИЕ"):
            sod = i
            break
    for p in d.paragraphs[rf + 1 : sod]:
        if not p.text.strip().startswith("Ключевые слова"):
            italize(p)
    d.save(fn)
    # length
    chars = sum(len(p.text) for p in d.paragraphs[rf + 1 : sod] if p.text.strip())
    return chars


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    c = process(fn)
    print(fn, "-> реферат символов:", c, f"(~{c//85} строк)")
