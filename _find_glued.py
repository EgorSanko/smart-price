# -*- coding: utf-8 -*-
"""Find glued/typo words across the diploma."""
import re
import io
from docx import Document

doc = Document(r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx")

# Collect all text with paragraph index
all_text = []
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        all_text.append((i, "P", p.text))

for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            for pi, p in enumerate(cell.paragraphs):
                if p.text.strip():
                    all_text.append((f"t{ti}r{ri}c{ci}p{pi}", "T", p.text))

# Patterns of problems
PATTERNS = [
    # Russian lowercase followed by Russian uppercase (word glue)
    (r"[а-яё][А-ЯЁ]", "mixed-case glue (рус)"),
    # Russian word + Latin word with no space
    (r"[а-яёА-ЯЁ][A-Za-z]", "cyr-lat glue"),
    (r"[A-Za-z][а-яёА-ЯЁ]", "lat-cyr glue"),
    # Word longer than 30 chars (likely glued)
    (r"\b\S{30,}\b", "extra-long token"),
    # Specific known issues
    (r"Россииипиа", "Россияиа"),
    (r"РоссиипиаБеларуси", "specific glue"),
    # Cyrillic letter + digit + Cyrillic letter — mid-word digit
    (r"[а-яё]\d+[а-яё]", "mid-word digit"),
]

# Whitelist (legit patterns)
WHITELIST_PATTERNS = [
    r"^[A-Za-z][A-Za-z]+\.(by|ru|com|org|net|io|me|app)$",  # domains
    r"iPhone|iPad|MacBook|YouTube|JavaScript|TypeScript|FastAPI|PostgreSQL",
    r"OpenRouter|OpenAI|GitHub|GitLab|Docker|Compose|EAlice",
    r"NextJS|NodeJS|React|Vue|Angular",
    r"^smrt-price\.ru$",
]


def is_whitelisted(token):
    for wp in WHITELIST_PATTERNS:
        if re.search(wp, token):
            return True
    return False


out = io.StringIO()
seen = set()

for idx, kind, text in all_text:
    for pat, label in PATTERNS:
        for m in re.finditer(pat, text):
            start = max(0, m.start() - 20)
            end = min(len(text), m.end() + 20)
            ctx = text[start:end]
            # extract surrounding token
            token_match = re.search(r"\S+", text[m.start() : m.end() + 30])
            token = token_match.group(0) if token_match else m.group(0)
            # extract full token: walk left to last whitespace, walk right to next whitespace
            ls = m.start()
            while ls > 0 and not text[ls - 1].isspace():
                ls -= 1
            le = m.end()
            while le < len(text) and not text[le].isspace():
                le += 1
            full_token = text[ls:le]

            if is_whitelisted(full_token):
                continue

            key = (idx, full_token, label)
            if key in seen:
                continue
            seen.add(key)
            out.write(f"[{idx}] [{label}] token={full_token!r}\n")
            out.write(f"    ctx: ...{ctx}...\n\n")

with open(r"C:/Users/egor3/Desktop/smart-price/_glued.txt", "w", encoding="utf-8") as f:
    f.write(out.getvalue())
print(f"Found {len(seen)} suspicious tokens. Written to _glued.txt")
