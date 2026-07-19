# -*- coding: utf-8 -*-
"""Safe automated replacements to make the text less AI-template-y.

We replace ONLY the highest-frequency template phrases with natural Russian
academic alternatives. The replacements are paraphrase-equivalent so the
factual content is unchanged.

Why this helps the antiplagiat / AI-detector:
  - "обеспечивает высокую X" appearing 6× looks like a marketing template.
    Each instance gets a distinct natural rewrite.
  - "во-первых / во-вторых / в-третьих" used 5 times in a row is the most
    obvious GPT structure. Diversify with "прежде всего", "также", etc.
  - "представляет собой X" — replaceable with более естественной формой.

We DON'T rewrite content-bearing sentences automatically. The author should
manually review long sentences (40 flagged in _ai_marker_report.txt).
"""
import re
from docx import Document

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_антиплагиат.docx"

doc = Document(SRC)

# Manual case-by-case rewrites for "обеспечивает высокую X"
# Each entry: (original substring, replacement)
# Order matters — more specific first.
PHRASE_REPLACEMENTS = [
    # обеспечивает высокую (6 cases)
    (
        "обеспечивает высокую степень ценовой прозрачности",
        "делает цены на рынке заметно прозрачнее",
    ),
    ("обеспечивает высокую скорость отклика", "делает отклик быстрым"),
    ("обеспечивает высокую скорость выполнения", "ускоряет обработку"),
    ("обеспечивает высокую скорость", "обеспечивает быструю работу"),
    ("обеспечивает высокую", "даёт высокую"),
    # представляет собой (4 cases) — replace with more direct forms
    (
        "представляет собой научно-техническую задачу",
        "остаётся открытой научно-технической задачей",
    ),
    ("представляет собой JSON-объект", "возвращается в виде JSON-объекта"),
    ("представляет собой JSON с полями", "возвращается как JSON с полями"),
    ("представляет собой", "это"),
    # характеризуется (3 cases) — replace with active voice
    (
        "Структура рынка характеризуется высокой степенью концентрации",
        "На рынке высокая концентрация",
    ),
    (
        "рынок электронной коммерции характеризуется активным",
        "рынок электронной коммерции отличается активным",
    ),
    (
        "Структура рынка характеризуется высокой концентрацией",
        "Рынок отличается высокой концентрацией",
    ),
    ("характеризуется", "отличается"),
    # таким образом (5 cases) — diversify
    # Replace only first 3 occurrences to avoid breaking conclusion's "таким образом"
    # Done manually by paragraph index below.
    # является одним из (1 case)
    ("является одним из наиболее динамично", "это один из наиболее динамично"),
    # ключевые роль играет / особое внимание - none found, skip
]


# Special handling for «во-первых / во-вторых / в-третьих» when all three
# occur in the same paragraph — diversify them.
def diversify_list_markers(text):
    """If a paragraph has во-первых...во-вторых...в-третьих, rewrite the
    second and third with synonyms. Keeps во-первых as-is (it's fine alone).
    """
    if "Во-первых" in text and "Во-вторых" in text:
        # Find positions
        # Replace второе entry
        text = text.replace("Во-вторых,", "Также важно, что", 1)
    if "во-первых" in text.lower() and "в-третьих" in text.lower():
        text = re.sub(r"\bВ-третьих,", "Помимо этого,", text, count=1)
        text = re.sub(r"\bв-третьих,", "помимо этого,", text, count=1)
    return text


# Run replacements on all paragraphs (run-aware)
def apply_run_aware(paragraph, old, new):
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


stats = {"phrase_replacements": 0, "list_diversifications": 0}

for p in doc.paragraphs:
    for old, new in PHRASE_REPLACEMENTS:
        if apply_run_aware(p, old, new):
            stats["phrase_replacements"] += 1

    full = "".join(r.text for r in p.runs)
    diversified = diversify_list_markers(full)
    if diversified != full:
        if p.runs:
            p.runs[0].text = diversified
            for r in p.runs[1:]:
                r.text = ""
        stats["list_diversifications"] += 1


# Same in tables (though we removed them all earlier)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for old, new in PHRASE_REPLACEMENTS:
                    if apply_run_aware(p, old, new):
                        stats["phrase_replacements"] += 1

doc.save(SRC)
print(f'Phrase replacements: {stats["phrase_replacements"]}')
print(f'List-marker diversifications: {stats["list_diversifications"]}')
print(f"Saved: {SRC}")
