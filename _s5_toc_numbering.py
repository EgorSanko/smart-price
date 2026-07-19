# -*- coding: utf-8 -*-
"""Script 5 — TOC single ПРИЛОЖЕНИЯ + page numbering (титул+реферат unnumbered).
"""
import copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def process(fn):
    d = Document(fn)
    paras = d.paragraphs

    # ---- 81. single ПРИЛОЖЕНИЯ in TOC ----
    appA = appB = None
    for p in paras:
        t = p.text.strip()
        if t.startswith("ПРИЛОЖЕНИЕ А") and p.style.name.startswith("Heading"):
            appA = p
        elif t.startswith("ПРИЛОЖЕНИЕ Б") and p.style.name.startswith("Heading"):
            appB = p
    if appA is not None:
        # insert "ПРИЛОЖЕНИЯ" Heading 1 before appendix А
        new = appA.insert_paragraph_before("ПРИЛОЖЕНИЯ", style="Heading 1")
        new.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new.paragraph_format.page_break_before = True
        for r in new.runs:
            r.font.size = Pt(14)
            r.bold = True
        # appendix А shares the divider page (no own break)
        appA.paragraph_format.page_break_before = False
        # exclude А and Б from TOC via direct outlineLvl 9
        for ap in (appA, appB):
            if ap is None:
                continue
            pPr = ap._element.get_or_add_pPr()
            ol = pPr.find(qn("w:outlineLvl"))
            if ol is None:
                ol = OxmlElement("w:outlineLvl")
                pPr.append(ol)
            ol.set(qn("w:val"), "9")

    # ---- 82. page numbering: section break before СОДЕРЖАНИЕ ----
    paras = d.paragraphs  # refresh after insert
    cont_idx = None
    for i, p in enumerate(paras):
        if p.text.strip() in ("СОДЕРЖАНИЕ", "ОГЛАВЛЕНИЕ"):
            cont_idx = i
            break
    if cont_idx is not None:
        cont = paras[cont_idx]
        prev = paras[cont_idx - 1]
        body = d.element.body
        body_sectPr = body.find(qn("w:sectPr"))
        new_sectPr = copy.deepcopy(body_sectPr)
        # section 0 = title+реферат: no footer/header -> no page number
        for tag in ("w:footerReference", "w:headerReference"):
            for e in new_sectPr.findall(qn(tag)):
                new_sectPr.remove(e)
        # type nextPage
        ty = new_sectPr.find(qn("w:type"))
        if ty is None:
            ty = OxmlElement("w:type")
            new_sectPr.insert(0, ty)
        ty.set(qn("w:val"), "nextPage")
        # no restart of numbering in section 0
        for pn in new_sectPr.findall(qn("w:pgNumType")):
            new_sectPr.remove(pn)
        # attach section break to the last paragraph of реферат (prev)
        prevPr = prev._element.get_or_add_pPr()
        # remove any existing sectPr to avoid duplicates
        for e in prevPr.findall(qn("w:sectPr")):
            prevPr.remove(e)
        prevPr.append(new_sectPr)
        # содержание: drop redundant page break (section break already new page)
        cont.paragraph_format.page_break_before = False
        # ensure final (section 1) numbering continues (no restart)
        for pn in body_sectPr.findall(qn("w:pgNumType")):
            body_sectPr.remove(pn)

    d.save(fn)
    return {"appendix_divider": appA is not None, "section_break_before": cont_idx}


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "->", process(fn))
