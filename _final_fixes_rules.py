# -*- coding: utf-8 -*-
"""Final rules-compliance fixes:
  1. Subsection headings (Heading 2, x.x) -> first-line indent 1.25 cm (rule 2.2).
  2. РЕФЕРАТ -> add «Объект исследования» / «Предмет исследования» (rule 1.3).
  3. Diagnostic: citation [N] placement (report only).
"""
import io
import re
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_нормоконтроль.docx"
doc = Document(SRC)
report = io.StringIO()

# 1. subsection indent 1.25 cm, left
sub = 0
for p in doc.paragraphs:
    if p.style.name == "Heading 2":
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(1.25)
        pf.left_indent = Cm(0)
        sub += 1

# 2. реферат объект/предмет — insert before "Целью"
inserted = False
for p in doc.paragraphs:
    t = p.text.strip()
    if (
        t.startswith("Целью данной работы")
        or t.startswith("Цель работы")
        or t.startswith("Целью работы")
    ):
        new = p.insert_paragraph_before(
            "Объект исследования — процессы поиска и сравнения цен на товары на "
            "маркетплейсах России и Беларуси. Предмет исследования — методы "
            "автоматизированного метапоиска и интеллектуального анализа цен с "
            "применением искусственного интеллекта."
        )
        # match реферат body formatting (Normal, justified, indent inherited)
        for r in new.runs:
            r.bold = False
            r.italic = False
        inserted = True
        break

# 3. citation placement diagnostic
cit = re.compile(r"\[\d+\]")
mid = 0
endish = 0
for p in doc.paragraphs:
    t = p.text
    for m in cit.finditer(t):
        after = t[m.end() : m.end() + 2].strip()
        # citation considered "at sentence end" if followed by . ; : or end
        if after[:1] in (".", ";", ":", ",", ")", ""):
            endish += 1
        else:
            mid += 1
report.write(f"Citations [N]: at end-ish={endish}, mid-sentence={mid}\n")

doc.save(SRC)
with open(
    r"C:/Users/egor3/Desktop/smart-price/_final_rules_report.txt", "w", encoding="utf-8"
) as f:
    f.write(report.getvalue())
print(f"Subsections indented: {sub}; реферат объект/предмет inserted: {inserted}")
print(report.getvalue())
