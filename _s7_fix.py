# -*- coding: utf-8 -*-
"""Script 7 — fix broken 'в таблице 10' reference + continuous title signature lines."""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NAMES = ("Санько", "Вагнер", "Чалова")
# pPr child order (subset) for correct pBdr placement
ORDER = [
    "pStyle",
    "keepNext",
    "keepLines",
    "pageBreakBefore",
    "framePr",
    "widowControl",
    "numPr",
    "suppressLineNumbers",
    "pBdr",
    "shd",
    "tabs",
    "suppressAutoHyphens",
    "kinsoku",
    "wordWrap",
    "overflowPunct",
    "topLinePunct",
    "autoSpaceDE",
    "autoSpaceDN",
    "bidi",
    "adjustRightInd",
    "snapToGrid",
    "spacing",
    "ind",
    "contextualSpacing",
    "mirrorIndents",
    "suppressOverlap",
    "jc",
    "textDirection",
    "textAlignment",
    "textboxTightWrap",
    "outlineLvl",
    "divId",
    "cnfStyle",
    "rPr",
    "sectPr",
    "pPrChange",
]


def insert_ordered(pPr, el):
    name = el.tag.split("}")[1]
    idx = ORDER.index(name)
    for child in pPr:
        cname = child.tag.split("}")[1]
        if cname in ORDER and ORDER.index(cname) > idx:
            child.addprevious(el)
            return
    pPr.append(el)


def add_bottom_border(p):
    pPr = p._element.get_or_add_pPr()
    old = pPr.find(qn("w:pBdr"))
    if old is not None:
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    insert_ordered(pPr, pBdr)


def process(fn):
    d = Document(fn)
    # 1. fix table 10 reference
    fixed_ref = False
    for p in d.paragraphs:
        t = p.text
        if "приведено в таблице" in t and "AI-моделей" in t and " 10" not in t:
            runs = p.runs
            for i, r in enumerate(runs):
                if r.text.rstrip().endswith("таблице"):
                    r.text = r.text.rstrip() + " 10."
                    # drop trailing whitespace/'.' runs after it
                    for rr in runs[i + 1 :]:
                        if rr.text.strip() in ("", "."):
                            rr.text = ""
                    fixed_ref = True
                    break
            break

    # 2. continuous title signature lines (bottom border instead of underline+tab)
    sig = 0
    for p in d.paragraphs[:30]:
        if any(n in p.text for n in NAMES):
            # remove char underline from runs
            for r in p.runs:
                rPr = r._element.find(qn("w:rPr"))
                if rPr is not None:
                    u = rPr.find(qn("w:u"))
                    if u is not None:
                        rPr.remove(u)
            # remove trailing tab run(s)
            for r in list(p.runs):
                if r.text == "\t":
                    r._element.getparent().remove(r._element)
            # remove tabs element
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                tabs = pPr.find(qn("w:tabs"))
                if tabs is not None:
                    pPr.remove(tabs)
            # add continuous bottom border
            add_bottom_border(p)
            sig += 1

    d.save(fn)
    return {"ref_fixed": fixed_ref, "signature_lines": sig}


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "->", process(fn))
