# -*- coding: utf-8 -*-
"""Unified normocontrol pass on a FRESH copy of master.

Per Latin token classification:
  - CODE  (snake_case, calls(), brackets[], curated camelCase) -> monospace
          font Consolas, NOT italic
  - TERM  (FastAPI, SSE, Docker, API, Gemini, ...)             -> italic

Plus:
  - Gemini 2.0 -> 2.5 (content sync with presentation)
  - Subheadings x.x / x.x.x: flush-left, indent 0, space_before 12pt,
    space_after 6pt, keep bold
  - Tables: single line spacing + same Latin treatment in cells
  - Реферат: пять -> шесть
"""
import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_нормоконтроль.docx"
doc = Document(SRC)

heading_re = re.compile(r"^\s*\d+(\.\d+){0,2}\s+\S")
subheading_re = re.compile(r"^\s*\d+\.\d+(\.\d+)?\s+\S")

TOKEN = re.compile(r"[A-Za-z_][\w]*(?:[.\-/][\w]+)*(?:\([^)]*\))?(?:\[[^\]]*\])?")

CURATED_CODE = {
    "asyncio.gather",
    "EventSourceResponse",
    "EAliceOfferCard",
    "EAliceOffer",
    "EAliceOfferList",
    "app.ml",
    "app.celery_app",
    "app.main",
    "sse_starlette",
}

stats = {"italic": 0, "mono": 0, "subheadings": 0, "cells": 0, "gemini": 0}


def classify(tok):
    if len(tok) < 2:
        return None
    if "_" in tok or "(" in tok or "[" in tok or tok in CURATED_CODE:
        return "code"
    if not re.search(r"[A-Za-z]", tok):
        return None
    return "term"


def segment(text):
    segs = []
    pos = 0
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            segs.append((text[pos : m.start()], None))
        segs.append((m.group(0), classify(m.group(0))))
        pos = m.end()
    if pos < len(text):
        segs.append((text[pos:], None))
    return segs


def run_is_simple(r_elem):
    for bad in ("w:br", "w:drawing", "w:pict", "w:tab", "w:fldChar", "w:object"):
        if r_elem.findall(".//" + qn(bad)):
            return False
    return len(r_elem.findall(qn("w:t"))) <= 1


def _rpr(r_elem):
    rPr = r_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = r_elem.makeelement(qn("w:rPr"), {})
        r_elem.insert(0, rPr)
    return rPr


def set_italic(r_elem):
    rPr = _rpr(r_elem)
    if rPr.find(qn("w:i")) is None:
        rPr.append(rPr.makeelement(qn("w:i"), {}))
    if rPr.find(qn("w:iCs")) is None:
        rPr.append(rPr.makeelement(qn("w:iCs"), {}))


def set_mono(r_elem):
    rPr = _rpr(r_elem)
    for tag in ("w:i", "w:iCs"):
        e = rPr.find(qn(tag))
        if e is not None:
            rPr.remove(e)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), "Consolas")


def process_paragraph(p):
    for r in list(p.runs):
        r_elem = r._element
        if not run_is_simple(r_elem):
            continue
        text = r.text
        if not text or not TOKEN.search(text):
            continue
        segs = segment(text)
        kinds = [k for _, k in segs if k]
        if not kinds:
            continue
        if len(segs) == 1:
            kind = segs[0][1]
            if kind == "term":
                set_italic(r_elem)
                stats["italic"] += 1
            elif kind == "code":
                set_mono(r_elem)
                stats["mono"] += 1
            continue
        parent = r_elem.getparent()
        idx = list(parent).index(r_elem)
        for off, (seg_text, kind) in enumerate(segs):
            new_r = deepcopy(r_elem)
            t = new_r.find(qn("w:t"))
            if t is None:
                t = new_r.makeelement(qn("w:t"), {})
                new_r.append(t)
            t.text = seg_text
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            if kind == "term":
                set_italic(new_r)
                stats["italic"] += 1
            elif kind == "code":
                set_mono(new_r)
                stats["mono"] += 1
            parent.insert(idx + off, new_r)
        parent.remove(r_elem)


intro_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "ВВЕДЕНИЕ":
        intro_idx = i
        break


def aware(p, old, new):
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return 0
    nf = full.replace(old, new)
    if p.runs:
        p.runs[0].text = nf
        for r in p.runs[1:]:
            r.text = ""
    return full.count(old)


GEMINI = [
    ("Gemini 2.0 Flash", "Gemini 2.5 Flash"),
    ("Gemini 2.0", "Gemini 2.5"),
    ("gemini-2.0-flash", "gemini-2.5-flash"),
    ("2.0 Flash", "2.5 Flash"),
]
REFERAT = [
    ("пять маркетплейсов", "шесть маркетплейсов"),
    ("объединяющая пять", "объединяющая шесть"),
    ("объединяет пять", "объединяет шесть"),
]

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    for o, nw in GEMINI:
        stats["gemini"] += aware(p, o, nw)
    if intro_idx is not None and i < intro_idx:
        for o, nw in REFERAT:
            aware(p, o, nw)

    if subheading_re.match(t) and intro_idx is not None and i > intro_idx:
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        stats["subheadings"] += 1
        continue
    if intro_idx is not None and i < intro_idx:
        continue
    if t.startswith("Ключевые слова"):
        continue
    if heading_re.match(t) or (t.isupper() and len(t) < 70):
        continue
    process_paragraph(p)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                stats["cells"] += 1
                if p.text.strip() and TOKEN.search(p.text):
                    process_paragraph(p)

doc.save(SRC)
print("=== STATS ===")
for k, v in stats.items():
    print(f"  {k}: {v}")
print("Saved.")
