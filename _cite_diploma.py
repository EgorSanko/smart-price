# -*- coding: utf-8 -*-
"""Scatter missing citations [11], [14], [17]-[25] into the body text.
Also replace source [11] Anthropic -> RFC 6455 WebSocket (actually used in cheaper feature).
"""
import sys, re
from docx import Document

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
doc = Document(SRC)


def insert_after_term_in_paragraph(paragraph, term, citation):
    """Insert ' [N]' right after the first whole-word occurrence of `term`.
    Works run-aware by rebuilding text if term spans runs."""
    text = paragraph.text
    if term not in text or citation in text:
        return False
    # try per-run
    for run in paragraph.runs:
        if term in run.text and citation not in run.text:
            run.text = run.text.replace(term, term + " " + citation, 1)
            return True
    # fallback: collapse
    new_text = text.replace(term, term + " " + citation, 1)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


# (paragraph_index, term_to_find_after, citation)
# indices from the earlier scan
insertions = [
    (171, "WebSocket", "[11]"),
    (373, "OpenRouter", "[14]"),  # section 3.7 has more direct OpenRouter context
    (91, "SQLAlchemy 2.0", "[17]"),
    (165, "Tailwind CSS", "[18]"),
    (153, "Playwright", "[19]"),
    (181, "PostgreSQL 16", "[20]"),
    (181, "Redis 7", "[21]"),
    (160, "Pydantic", "[22]"),
    (165, "Zustand", "[23]"),
    (165, "TanStack React Query", "[24]"),
    (165, "Recharts", "[25]"),
]

ok = 0
for idx, term, cite in insertions:
    p = doc.paragraphs[idx]
    # if paragraph doesn't contain term (shifted after earlier edits), search nearby
    if term not in p.text:
        found = False
        for j in range(max(0, idx - 3), min(len(doc.paragraphs), idx + 4)):
            if term in doc.paragraphs[j].text:
                p = doc.paragraphs[j]
                found = True
                break
        if not found:
            print(f"MISS: {term} {cite} near {idx}")
            continue
    if insert_after_term_in_paragraph(p, term, cite):
        print(f'OK  [{cite}] after "{term}"')
        ok += 1
    else:
        print(f'SKIP [{cite}] after "{term}" (already cited or not found)')

# --- Replace source [11] in bibliography ---
new_src11 = (
    "11. Fette, I. The WebSocket Protocol / I. Fette, A. Melnikov. — "
    "IETF, 2011. — RFC 6455. — Режим доступа: "
    "https://datatracker.ietf.org/doc/html/rfc6455. — Дата доступа: 10.04.2026."
)
for p in doc.paragraphs:
    if p.text.lstrip().startswith("11. Anthropic"):
        # rewrite paragraph text preserving first-run formatting
        if p.runs:
            p.runs[0].text = new_src11
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new_src11)
        print("Replaced source [11]")
        break

doc.save(SRC)
print(f"Done. Citations inserted: {ok}/{len(insertions)}")
