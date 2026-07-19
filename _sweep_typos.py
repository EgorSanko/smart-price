# -*- coding: utf-8 -*-
"""Comprehensive typo/glue sweep across entire diploma — paragraphs + tables + headers/footers."""
import re
import io
from docx import Document

doc = Document(r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx")


# Collect every text fragment in the document
def all_text_with_loc(d):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip():
            yield (f"P{i}", p.text)
    for ti, t in enumerate(d.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    if p.text.strip():
                        yield (f"T{ti}r{ri}c{ci}p{pi}", p.text)
    # headers/footers
    for si, sec in enumerate(d.sections):
        for hf, name in [(sec.header, "header"), (sec.footer, "footer")]:
            if hf:
                for pi, p in enumerate(hf.paragraphs):
                    if p.text.strip():
                        yield (f"S{si}.{name}.p{pi}", p.text)


# Common Russian "particles" that often get glued to neighbors
# When a paragraph has a 2-3 letter Russian word stuck to a longer word, it's likely glue
SHORT_PARTICLES = {
    "и",
    "в",
    "на",
    "с",
    "а",
    "но",
    "или",
    "из",
    "к",
    "о",
    "у",
    "до",
    "по",
    "за",
    "для",
    "над",
    "под",
}


def find_problems(text):
    issues = []

    # 1. Lowercase letter directly followed by uppercase letter inside a word (likely glue)
    for m in re.finditer(r"\b\w*[а-яё][А-ЯЁ]\w*\b", text):
        word = m.group(0)
        # exclude legit proper nouns where this is normal — none in Russian, so flag all
        issues.append(("case-glue", word, m.start()))

    # 2. Letter immediately followed by a digit followed by letter (mid-word digit)
    for m in re.finditer(r"\b\w*[А-ЯЁа-яё]\d+[А-ЯЁа-яё]\w*\b", text):
        word = m.group(0)
        # exclude common like "5G", "Gemini2" etc. — but those wouldn't have cyrillic surrounding
        issues.append(("digit-glue", word, m.start()))

    # 3. Cyrillic word followed by Latin word with no space — sometimes legit (e.g. "API Gemini")
    #    but flag mid-word cases
    for m in re.finditer(r"\b[А-ЯЁа-яё]+[A-Za-z]+\b", text):
        word = m.group(0)
        issues.append(("cyr-lat-glue", word, m.start()))

    # 4. Latin followed by Cyrillic mid-word
    for m in re.finditer(r"\b[A-Za-z]+[А-ЯЁа-яё]+\b", text):
        word = m.group(0)
        # exclude very common false positives
        issues.append(("lat-cyr-glue", word, m.start()))

    # 5. Words with double consecutive vowel patterns that seem unnatural ("ииа", "ееи", "оои" etc.)
    UNNATURAL_TRIPLES = ["ииа", "иио", "иие", "ыыа", "ееи", "еее", "ооо", "ыыи", "аааа"]
    for m in re.finditer(r"\b\w*([а-яё])\1\w+", text):
        word = m.group(0)
        # only triples — already implicit ([a-zа-я])\1 means doubled
        # let's check explicit unnatural
        for pat in UNNATURAL_TRIPLES:
            if pat in word.lower():
                issues.append(("unnatural-triple", word, m.start()))
                break

    # 6. Word longer than 25 cyrillic chars
    for m in re.finditer(r"\b[А-ЯЁа-яё-]{25,}\b", text):
        word = m.group(0)
        issues.append(("long-word", word, m.start()))

    return issues


# Whitelist for known-good multi-script tokens
WHITELIST = {
    # English brand names
    "iPhone",
    "iPad",
    "iPod",
    "Mac",
    "macOS",
    "iOS",
    # Common product names with Russian
    "Onliner",
    "OpenRouter",
    "OpenAI",
    "JavaScript",
    "TypeScript",
}


def is_whitelisted(token):
    # whitelist exact match or via common prefix
    if token in WHITELIST:
        return True
    # numbers like "5G", "2.5"
    if re.fullmatch(r"\d+[A-Za-z]+", token):
        return True
    return False


out = io.StringIO()
all_issues = {}  # type: dict[(loc, kind, word) -> ctx]

for loc, text in all_text_with_loc(doc):
    issues = find_problems(text)
    for kind, word, start in issues:
        if is_whitelisted(word):
            continue
        ctx_start = max(0, start - 30)
        ctx_end = min(len(text), start + len(word) + 30)
        ctx = text[ctx_start:ctx_end]
        key = (loc, kind, word)
        if key in all_issues:
            continue
        all_issues[key] = ctx

# Dedupe: same word may appear in many places — group by word
from collections import defaultdict

by_word = defaultdict(list)
for (loc, kind, word), ctx in all_issues.items():
    by_word[(kind, word)].append((loc, ctx))

# Sort by kind, then word
order = sorted(by_word.items(), key=lambda x: (x[0][0], x[0][1]))

for (kind, word), locs in order:
    out.write(f"[{kind}] {word!r} (×{len(locs)})\n")
    for loc, ctx in locs[:3]:
        out.write(f"  @ {loc}: ...{ctx}...\n")
    if len(locs) > 3:
        out.write(f"  ... and {len(locs) - 3} more\n")
    out.write("\n")

with open(r"C:/Users/egor3/Desktop/smart-price/_sweep.txt", "w", encoding="utf-8") as f:
    f.write(out.getvalue())
print(
    f"Total unique issues: {len(by_word)}, total occurrences: {sum(len(v) for v in by_word.values())}"
)
