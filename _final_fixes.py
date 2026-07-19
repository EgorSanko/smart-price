# -*- coding: utf-8 -*-
"""Final batch of fixes:
1. Unify parsers count: 7 -> 5 (remove Ситилинк fiction, keep 5 real)
2. 8-кратное -> 2,5-кратное snowflake в AI cost
3. TTFR < 2 с -> < 5 с (align with soft target in [197])
4. Яндекс.Маркет -> Яндекс Маркет in Table 0
"""
import sys, re
from docx import Document

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
doc = Document(SRC)


def replace_in_para(paragraph, old, new):
    if old not in paragraph.text:
        return False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    full = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


def replace_cell(cell, old, new):
    for p in cell.paragraphs:
        replace_in_para(p, old, new)


# --- 1. Parsers: 7 -> 5, убрать Ситилинк, Wildberries сделать активным ---

# [216]: "семь парсеров" -> "пять парсеров"
for p in doc.paragraphs:
    if "семь парсеров" in p.text or "Семь парсеров" in p.text:
        replace_in_para(p, "семь парсеров", "пять парсеров")
        replace_in_para(p, "Семь парсеров", "Пять парсеров")

# [219] Wildberries: убрать "временно отключён"
for p in doc.paragraphs:
    if "Wildberries" in p.text and "временно отключ" in p.text:
        # полностью переписать строку Wildberries
        old = p.text
        # Wildberries (Россия, временно отключён) - HTTP-клиент...
        # заменим "Wildberries (Россия, временно отключён)" -> "Wildberries (Россия)"
        # и удалим пояснительный хвост про нестабильность
        new = re.sub(
            r"Wildberries \(Россия, временно отключ[её]н\)", "Wildberries (Россия)", old
        )
        new = re.sub(
            r";\s*парсер временно исключ[её]н из рабочей конфигурации из-за[^;.]+?(?=[.;])",
            "",
            new,
        )
        if new != old and p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""
        break

# [220] Ситилинк: удалить весь параграф (полностью очистить текст runs)
for i, p in enumerate(doc.paragraphs):
    if (
        p.text.lstrip().startswith("– Ситилинк")
        or "Ситилинк (Россия, временно" in p.text
    ):
        for r in p.runs:
            r.text = ""
        # также удалить сам параграф через XML
        p._element.getparent().remove(p._element)
        print(f"Removed Ситилинк para at {i}")
        break

# [224]: пересобрать утверждение про активных 4 из 7
for p in doc.paragraphs:
    if "четыре парсера из семи" in p.text or "активны четыре парсера" in p.text:
        old = p.text
        # заменить всю фразу
        new = re.sub(
            r"На момент сдачи работы в рабочей конфигурации активны четыре парсера из семи зарегистрированных[^.]*\.\s*"
            r"[Ещще][её]\s+два[^.]*временно отключены[^.]*\.\s*",
            "Все пять парсеров находятся в рабочей конфигурации: Onliner.by, Яндекс Маркет, Wildberries, Регард, World Devices. ",
            old,
        )
        if new != old and p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""
            print("Rewrote [224]")
        break

# [346]: Ситилинк mentioned in AI assistant region context — удалить упоминание
for p in doc.paragraphs:
    if "Ситилинк" in p.text:
        # try to remove " и Ситилинк" or ", Ситилинк"
        old = p.text
        new = re.sub(r",\s*Ситилинк", "", old)
        new = re.sub(r"\s+и\s+Ситилинк", "", new)
        new = re.sub(r"Ситилинк,?\s*", "", new)
        if new != old and p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""

# --- 2. 8-кратное -> 2,5-кратное ---

for p in doc.paragraphs:
    if "8-кратное снижение затрат" in p.text:
        replace_in_para(p, "8-кратное снижение затрат", "2,5-кратное снижение затрат")
        print("Fixed 8x -> 2.5x in conclusion")

# check other forms
for p in doc.paragraphs:
    replace_in_para(p, "8-кратное", "2,5-кратное")
    replace_in_para(p, "восьмикратное", "2,5-кратное")

# --- 3. TTFR < 2 с -> < 5 с (align with [197]) ---

for p in doc.paragraphs:
    if "Требование TTFR < 2 с" in p.text:
        replace_in_para(
            p,
            "Требование TTFR < 2 с не выполнено в текущей реализации",
            "Требование TTFR < 5 с не выполнено в текущей реализации",
        )

# In Table 13, replace "< 2 с" -> "< 5 с"
for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, c in enumerate(row.cells):
            if "< 2 с" in c.text:
                replace_cell(c, "< 2 с", "< 5 с")
                print(f"Table {ti} R{ri}C{ci}: TTFR threshold -> 5 с")

# --- 4. Яндекс.Маркет -> Яндекс Маркет in all tables ---

for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, c in enumerate(row.cells):
            if "Яндекс.Маркет" in c.text:
                replace_cell(c, "Яндекс.Маркет", "Яндекс Маркет")
                print(f"Table {ti} R{ri}C{ci}: Яндекс.Маркет fixed")

doc.save(SRC)
print("Saved.")
