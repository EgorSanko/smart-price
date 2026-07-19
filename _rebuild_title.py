# -*- coding: utf-8 -*-
"""Rebuild Sanko's title page by cloning Ишимова's СПбГУПТД template and
substituting Sanko's data. Guarantees identical official formatting.
"""
import copy
from docx import Document

ISH = r"C:/Users/egor3/Downloads/Ишимова ВКР-6.docx"
SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_нормоконтроль.docx"

ish = Document(ISH)

# Find РЕФЕРАТ in Ишимова -> title block is everything before it
ish_ref = next(i for i, p in enumerate(ish.paragraphs) if p.text.strip() == "РЕФЕРАТ")
title_paras = ish.paragraphs[:ish_ref]

REPLACE = [
    (
        "10.03.01, Информационная безопасность",
        "09.03.02 – Информационные системы и технологии (Информационные системы и сетевые технологии)",
    ),
    (
        "Анализ и модернизация системы безопасности компьютерной сети предприятия",
        "Система метапоиска товаров на торговых площадках с функцией интеллектуального анализа цен",
    ),
    ("4-МД-9", "4-МД-16"),
    ("Ишимова Дарья Антоновна", "Санько Егор Александрович"),
    # руководитель: у Санько — Вагнер (степень неизвестна, оставляем ФИО)
    (
        "к.т.н.  доцент кафедры интеллектуальных систем и защиты информации, Чалова Екатерина Игорьевна",
        "Вагнер Виктория Игоревна",
    ),
    (
        "к.т.н. доцент кафедры интеллектуальных систем и защиты информации, Чалова Екатерина Игорьевна",
        "Вагнер Виктория Игоревна",
    ),
    # нормоконтролёр: Чалова Е.И. (как у Ишимовой и Гусенова) — оставляем без замены
]


def aware(p, old, new):
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    nf = full.replace(old, new)
    if p.runs:
        p.runs[0].text = nf
        for r in p.runs[1:]:
            r.text = ""
    return True


# apply replacements; mark консультанты paragraph for skip
clone_elems = []
for p in title_paras:
    if "Консультант" in p.text:
        continue  # drop consultants line (Sanko has none)
    for old, new in REPLACE:
        aware(p, old, new)
    clone_elems.append(copy.deepcopy(p._p))

# --- Sanko ---
sanko = Document(SRC)
s_ref_para = next(p for p in sanko.paragraphs if p.text.strip() == "РЕФЕРАТ")
s_ref_el = s_ref_para._p

# delete Sanko's current title paragraphs (everything before РЕФЕРАТ)
to_del = []
for p in sanko.paragraphs:
    if p._p is s_ref_el:
        break
    to_del.append(p._p)
for el in to_del:
    el.getparent().remove(el)

# insert cloned title elements before РЕФЕРАТ
for el in clone_elems:
    s_ref_el.addprevious(el)

# ensure РЕФЕРАТ starts on a new page
s_ref_para.paragraph_format.page_break_before = True

sanko.save(SRC)
print(
    f"Title rebuilt: {len(clone_elems)} paragraphs cloned, {len(to_del)} old removed."
)
