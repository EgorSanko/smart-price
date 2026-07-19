# -*- coding: utf-8 -*-
"""Apply edits to ДИПЛОМ.docx in place.

Changes:
1. Unify 'Яндекс.Маркет' -> 'Яндекс Маркет'
2. Abstract: '22 рисунка' -> '18 рисунков'
3. Renumber figures 2.2..2.6 -> 2.1..2.5 (fills the gap)
4. Insert new TOC entry '3.13 Реализация функции «Найти дешевле»'
5. Rename existing '3.13 Выводы по третьей главе' -> '3.14 ...' (TOC + body)
6. Insert new section body before current 3.13 heading
7. Extend conclusions paragraph to mention cheaper feature
"""
import sys, re, copy
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
doc = Document(SRC)

# ---- helpers --------------------------------------------------------------


def replace_in_runs(paragraph, old, new):
    """Safer than editing paragraph.text: preserves formatting when possible.
    Falls back to writing into first run if string is split across runs."""
    if old not in paragraph.text:
        return False
    # try per-run
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # split across runs: dump to first run, clear others
    full = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


def insert_paragraph_before(paragraph, text, copy_format_from=None):
    """Insert a paragraph before given paragraph; copy formatting from template para."""
    template = copy_format_from if copy_format_from is not None else paragraph
    new_el = copy.deepcopy(template._element)
    # strip existing runs
    for r in new_el.findall(qn("w:r")):
        new_el.remove(r)
    for hl in new_el.findall(qn("w:hyperlink")):
        new_el.remove(hl)
    paragraph._element.addprevious(new_el)
    # wrap into Paragraph via doc.paragraphs is messy; use python-docx internal
    from docx.text.paragraph import Paragraph

    new_p = Paragraph(new_el, paragraph._parent)
    # copy first run formatting from template if any
    if template.runs:
        src_run = template.runs[0]
        run = new_p.add_run(text)
        try:
            run.font.name = src_run.font.name
            if src_run.font.size:
                run.font.size = src_run.font.size
            run.bold = src_run.bold
            run.italic = src_run.italic
        except Exception:
            pass
    else:
        new_p.add_run(text)
    return new_p


# ---- 1. Unify Яндекс.Маркет -> Яндекс Маркет ------------------------------

replaced_ya = 0
for p in doc.paragraphs:
    if "Яндекс.Маркет" in p.text:
        replace_in_runs(p, "Яндекс.Маркет", "Яндекс Маркет")
        replaced_ya += 1

# ---- 2. Abstract figure count 22 -> 18 ------------------------------------

for p in doc.paragraphs:
    if "22 рисунка" in p.text:
        replace_in_runs(p, "22 рисунка", "18 рисунков")
        break

# ---- 3. Renumber figures 2.2..2.6 -> 2.1..2.5 -----------------------------

# Do in descending order to avoid collisions not needed since targets all distinct
fig_renames = [
    ("Рисунок 2.2", "Рисунок 2.1"),
    ("Рисунок 2.3", "Рисунок 2.2"),
    ("Рисунок 2.4", "Рисунок 2.3"),
    ("Рисунок 2.5", "Рисунок 2.4"),
    ("Рисунок 2.6", "Рисунок 2.5"),
    # also references 'рисунке 2.X' / 'рисунок 2.X' lowercase
    ("рисунке 2.2", "рисунке 2.1"),
    ("рисунке 2.3", "рисунке 2.2"),
    ("рисунке 2.4", "рисунке 2.3"),
    ("рисунке 2.5", "рисунке 2.4"),
    ("рисунке 2.6", "рисунке 2.5"),
    ("рисунок 2.2", "рисунок 2.1"),
    ("рисунок 2.3", "рисунок 2.2"),
    ("рисунок 2.4", "рисунок 2.3"),
    ("рисунок 2.5", "рисунок 2.4"),
    ("рисунок 2.6", "рисунок 2.5"),
]
# Use a two-pass approach with unique sentinel to avoid chain collisions
SENT = "\u2603"
for p in doc.paragraphs:
    for old, new in fig_renames:
        if old in p.text:
            sentinel = new.replace("2.", f"2.{SENT}")
            replace_in_runs(p, old, sentinel)
for p in doc.paragraphs:
    if SENT in p.text:
        replace_in_runs(p, SENT, "")

# ---- 4/5/6. TOC update + body insert + rename 3.13->3.14 -------------------

# Find TOC entry for existing 3.13 and the body heading
toc_3_13_idx = None
body_3_13_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("3.13 Выводы по третьей главе"):
        if toc_3_13_idx is None:
            toc_3_13_idx = i
        else:
            body_3_13_idx = i
            break

assert toc_3_13_idx is not None and body_3_13_idx is not None, (
    toc_3_13_idx,
    body_3_13_idx,
)
print(f"TOC 3.13 at {toc_3_13_idx}, body 3.13 at {body_3_13_idx}")

toc_para = doc.paragraphs[toc_3_13_idx]
body_para = doc.paragraphs[body_3_13_idx]

# Insert TOC entry for new 3.13 before existing TOC line
new_toc_text = "3.13 Реализация функции «Найти дешевле» ................... —"
insert_paragraph_before(toc_para, new_toc_text, copy_format_from=toc_para)

# Rename existing TOC line 3.13 -> 3.14
replace_in_runs(
    toc_para, "3.13 Выводы по третьей главе", "3.14 Выводы по третьей главе"
)

# Body: insert new section before body 3.13 heading
new_body = [
    "3.13 Реализация функции поиска более выгодных предложений («Найти дешевле»)",
    "3.13.1 Назначение подсистемы",
    "Функция «Найти дешевле» реализует сценарий повторного поиска для конкретной товарной позиции: пользователь указывает прямую ссылку на карточку товара маркетплейса Яндекс Маркет, после чего система в автоматическом режиме подбирает у сторонних продавцов предложения идентичного товара по меньшей цене. Функция дополняет основной сценарий потокового поиска и ориентирована на случаи, когда пользователь уже выбрал конкретную модель и нуждается в проверке, не является ли текущее предложение избыточно дорогим. Подсистема реализована в виде асинхронной задачи Celery, выполняющей длительное взаимодействие с внешним источником данных без блокировки HTTP-запроса.",
    "3.13.2 Источник данных и протокол взаимодействия",
    "В качестве источника данных используется внутренний канал подбора ценовых предложений, применяемый голосовым помощником Яндекс. Доступ реализован через постоянное соединение по протоколу WebSocket с публичным шлюзом. Клиент, реализованный в модуле workers/alisa.py, выполняет последовательность шагов: установление защищённого соединения, прохождение процедуры инициализации сессии, отправка текстового запроса на поиск более дешёвых предложений для заданного идентификатора товара (SKU), приём и декодирование потока сообщений rich_uicard и нормальное закрытие соединения. Каждое сообщение передаётся в двоичном формате и содержит вложенное JSON-поле json_data с типизированной структурой EAliceOfferCard, EAliceOffer или EAliceOfferList.",
    "3.13.3 Алгоритм извлечения предложений",
    "Ключевой особенностью обработки ответа является извлечение всех предложений из объекта типа EAliceOfferList, агрегирующего до нескольких десятков альтернативных продавцов в едином сообщении. В первоначальной реализации парсер ограничивался обработкой верхнеуровневой структуры, что приводило к потере значительной части данных — из четырнадцати доступных предложений обрабатывались только три-четыре. Для устранения этой проблемы функция обхода _iter_json_data_payloads расширена рекурсивным спуском в массив offerList, при этом сохранена поддержка одиночных карточек EAliceOfferCard и EAliceOffer. Итоговое покрытие возросло более чем втрое и составило в среднем одиннадцать-четырнадцать предложений на запрос.",
    "3.13.4 Обнаружение отказов и завершение задачи",
    "Поскольку удалённый сервис способен возвращать текстовый отказ («пока не умею искать такие товары», «не поддерживается» и иные формулировки), подсистема реализует детектор отказов на основе регулярного выражения. При совпадении ответа с любым из шаблонов задача завершается со статусом rejected в среднем за пятнадцать-двадцать секунд, что предотвращает бессмысленное ожидание до таймаута в десять минут. Список формулировок отказа сформирован эмпирически по результатам наблюдения за реальными ответами сервиса для нетипичных запросов (напитки, кондитерские изделия, товары ограниченного оборота). Статус задачи публикуется по мере обработки через отдельное SSE-соединение и принимает значения idle, searching, complete, rejected, что позволяет клиенту корректно отображать прогресс.",
    "3.13.5 Пользовательский интерфейс",
    "Подсистема снабжена отдельной страницей /cheaper с набором состояний, визуально согласованных с графическим стилем проекта. Для каждого состояния задействована собственная векторная анимированная иллюстрация: копилка со скатывающимися монетами — для исходного ожидания ввода; катящаяся корзина — для активной фазы поиска; улыбающаяся монета — для успешного завершения с указанием найденной экономии; грустная панда — для отказа сервиса. Вспомогательная плашка в верхней части страницы информирует об ожидаемой продолжительности операции («примерно десять минут»), а нижняя плашка отображает реальное прошедшее время и сообщение о возможности закрыть страницу без потери результата. Таким образом, пользовательский опыт приведён в соответствие с особенностями длительного асинхронного сценария.",
]
# Use the Выводы heading paragraph as formatting template for headings;
# the paragraph immediately after (body text) as template for body text.
heading_tmpl = body_para
# text template — one paragraph after heading
text_tmpl = (
    doc.paragraphs[body_3_13_idx + 1]
    if body_3_13_idx + 1 < len(doc.paragraphs)
    else body_para
)

for line in new_body:
    is_heading = bool(re.match(r"^3\.13(\.\d+)?\s", line))
    tmpl = heading_tmpl if is_heading else text_tmpl
    insert_paragraph_before(body_para, line, copy_format_from=tmpl)

# Rename body heading 3.13 -> 3.14
replace_in_runs(
    body_para, "3.13 Выводы по третьей главе", "3.14 Выводы по третьей главе"
)

# ---- 7. Extend conclusions text with cheaper mention -----------------------
# find paragraph that immediately follows '3.14 Выводы по третьей главе'
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("3.14 Выводы по третьей главе"):
        # append to next substantive paragraph
        nxt = doc.paragraphs[i + 1]
        add = " Реализована функция «Найти дешевле», позволяющая по ссылке на товар Яндекс Маркета подобрать более дешёвые предложения у сторонних продавцов через асинхронную Celery-задачу с извлечением данных из вложенной структуры EAliceOfferList и детектированием отказов по регулярному выражению."
        if nxt.runs:
            nxt.runs[-1].text = nxt.runs[-1].text.rstrip() + add
        break

# ---- save ------------------------------------------------------------------
OUT = SRC
doc.save(OUT)
print(f"Done. Replaced Яндекс.Маркет: {replaced_ya}. Saved to {OUT}")
