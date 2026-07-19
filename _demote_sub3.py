# -*- coding: utf-8 -*-
"""п.4 — Demote all x.x.x subheadings to bold run-in lead-ins.

Each "x.x.x Название" heading line is removed; "Название. " is prepended
as a BOLD run to the following body paragraph. Content preserved, the
third structural level disappears (doc keeps chapter + x.x only).
TOC lists only x.x, so no TOC cleanup needed.
"""
import re
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_нормоконтроль.docx"
doc = Document(SRC)

sub3 = re.compile(r"^\s*\d+\.\d+\.\d+\s+(.+?)\s*$")


def del_paragraph(p):
    p._element.getparent().remove(p._element)


def prepend_bold(target_p, text):
    new_r = target_p.add_run(text)
    new_r.bold = True
    r_el = new_r._element
    target_p._p.remove(r_el)
    ppr = target_p._p.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(r_el)
    else:
        target_p._p.insert(0, r_el)


paras = doc.paragraphs
jobs = []
for i, p in enumerate(paras):
    m = sub3.match(p.text.strip())
    if not m:
        continue
    title = m.group(1).rstrip(".")
    nxt = None
    for j in range(i + 1, len(paras)):
        if paras[j].text.strip():
            nxt = paras[j]
            break
    jobs.append((p, nxt, title))

done = 0
for heading_p, nxt_p, title in jobs:
    if nxt_p is None:
        continue
    prepend_bold(nxt_p, title + ". ")
    del_paragraph(heading_p)
    done += 1

doc.save(SRC)
print(f"Demoted x.x.x subheadings: {done} / {len(jobs)}")
