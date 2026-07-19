# -*- coding: utf-8 -*-
"""Inspect body paragraph spacing + chapter page-breaks + abstract length + Latin token sample."""
import io
import re
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
doc = Document(SRC)
out = io.StringIO()

heading_re = re.compile(r"^\s*\d+(\.\d+){0,2}\s+\S")

# 1. Body paragraph spacing sample (first 10 real body paragraphs after ВВЕДЕНИЕ)
out.write("=== BODY PARAGRAPH SPACING (after ВВЕДЕНИЕ) ===\n")
started = False
shown = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if t == "ВВЕДЕНИЕ":
        started = True
        continue
    if not started or not t:
        continue
    if heading_re.match(t):
        continue
    pf = p.paragraph_format
    out.write(
        f"  first_indent={pf.first_line_indent} space_before={pf.space_before} "
        f"space_after={pf.space_after} line_spacing={pf.line_spacing} "
        f"align={pf.alignment}\n"
    )
    out.write(f"    {t[:60]!r}\n")
    shown += 1
    if shown >= 6:
        break

# 2. Chapter heading page-break check
out.write("\n=== CHAPTER PAGE-BREAK (level-1 headings) ===\n")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r"^[1-4]\s+[А-ЯЁ]{4,}", t):
        # check page_break_before
        pf = p.paragraph_format
        pbb = pf.page_break_before
        # check for w:br type=page in runs
        has_br = any(r._element.findall(".//" + qn("w:br")) for r in p.runs)
        out.write(f"  [{i}] page_break_before={pbb} text={t[:50]!r}\n")

# 3. Abstract (РЕФЕРАТ) length
out.write("\n=== РЕФЕРАТ length ===\n")
ab_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "РЕФЕРАТ":
        ab_start = i
        break
if ab_start is not None:
    chars = 0
    paras = 0
    for p in doc.paragraphs[ab_start + 1 :]:
        t = p.text.strip()
        if t == "СОДЕРЖАНИЕ":
            break
        if t:
            chars += len(t)
            paras += 1
    out.write(
        f"  Paragraphs: {paras}, chars: {chars}, approx lines (90ch): {chars//90}\n"
    )
    out.write("  --- abstract text ---\n")
    for p in doc.paragraphs[ab_start + 1 : ab_start + 1 + paras + 1]:
        t = p.text.strip()
        if t == "СОДЕРЖАНИЕ":
            break
        if t:
            out.write(f"  | {t}\n")

# 4. Latin token sample — how many runs/paragraphs contain Latin
out.write("\n=== LATIN TOKEN PREVALENCE ===\n")
latin_re = re.compile(r"[A-Za-z][A-Za-z0-9]*(?:[.\-+][A-Za-z0-9]+)*")
paras_with_latin = 0
sample_tokens = set()
for p in doc.paragraphs:
    toks = latin_re.findall(p.text)
    toks = [x for x in toks if len(x) >= 2]
    if toks:
        paras_with_latin += 1
        for x in toks[:50]:
            if len(sample_tokens) < 60:
                sample_tokens.add(x)
out.write(f"  Paragraphs containing Latin tokens (>=2ch): {paras_with_latin}\n")
out.write(f"  Sample tokens: {sorted(sample_tokens)}\n")

with open(
    r"C:/Users/egor3/Desktop/smart-price/_format_inspect3.txt", "w", encoding="utf-8"
) as f:
    f.write(out.getvalue())
print("OK inspect3")
