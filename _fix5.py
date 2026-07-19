# -*- coding: utf-8 -*-
"""Five-point fix:
  A. Keywords -> lowercase (Russian), Latin keeps proper case (+italic later).
  C. Tables: rows cantSplit (no mid-row page break).
  E. Un-bold body lead-ins (demoted x.x.x + выводы) — keep only headings bold.
  B. TOC styles: right dot-leader tab at 16.5 cm so page numbers align.
"""
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ_нормоконтроль.docx"
doc = Document(SRC)
stats = {"keywords": 0, "cantsplit_rows": 0, "unbold": 0, "toc_tabs": 0}

# boundaries
intro = referat = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "РЕФЕРАТ" and referat is None:
        referat = i
    if t == "ВВЕДЕНИЕ" and intro is None:
        intro = i

# ---------- A. keywords lowercase ----------
NEW_KW = (
    "Ключевые слова: метапоиск, парсинг, маркетплейс, искусственный интеллект, "
    "сравнение цен, скрапинг, веб-приложение, FastAPI, Next.js, Gemini, OpenRouter, "
    "SSE, потоковая передача данных, история цен, агрегатор цен, электронная коммерция, "
    "Docker, PostgreSQL, антибот-защита."
)
for p in doc.paragraphs[: referat + 30 if referat else 60]:
    if p.text.strip().startswith("Ключевые слова"):
        if p.runs:
            p.runs[0].text = NEW_KW
            for r in p.runs[1:]:
                r.text = ""
        stats["keywords"] += 1
        break

# ---------- C. tables: rows cantSplit ----------
for table in doc.tables:
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))
            stats["cantsplit_rows"] += 1

# ---------- E. un-bold body lead-ins ----------
for i, p in enumerate(doc.paragraphs):
    if intro is None or i <= intro:
        continue
    if p.style.name != "Normal":
        continue  # keep Heading 1/2 bold
    if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        continue  # spare any centered structural text
    for r in p.runs:
        # un-bold (keep italic / mono)
        rPr = r._element.find(qn("w:rPr"))
        if rPr is not None:
            b = rPr.find(qn("w:b"))
            if b is not None:
                rPr.remove(b)
                stats["unbold"] += 1
            bcs = rPr.find(qn("w:bCs"))
            if bcs is not None:
                rPr.remove(bcs)

# ---------- B. TOC styles: right dot tab at 16.5 cm ----------
for nm in ["toc 1", "toc 2"]:
    try:
        st = doc.styles[nm]
    except KeyError:
        continue
    pf = st.paragraph_format
    # clear existing tabs
    pPr = st.element.get_or_add_pPr()
    old = pPr.find(qn("w:tabs"))
    if old is not None:
        pPr.remove(old)
    pf.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    pf.left_indent = Cm(0)
    pf.first_line_indent = Cm(0)
    stats["toc_tabs"] += 1

# re-assert updateFields
s = doc.settings.element
for e in s.findall(qn("w:updateFields")):
    s.remove(e)
uf = OxmlElement("w:updateFields")
uf.set(qn("w:val"), "true")
s.insert(0, uf)

doc.save(SRC)
print(stats)
