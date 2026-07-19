# -*- coding: utf-8 -*-
"""Script 2 — compact РЕФЕРАТ to one page + keywords CAPS."""
from docx import Document
from docx.oxml.ns import qn

NEW = [
    "Выпускная квалификационная работа содержит 99 страниц, 18 рисунков, "
    "23 таблицы, 25 использованных источников, 2 приложения.",
    "КЛЮЧЕВЫЕ СЛОВА: МЕТАПОИСК, ПАРСИНГ, МАРКЕТПЛЕЙС, ИСКУССТВЕННЫЙ ИНТЕЛЛЕКТ, "
    "СРАВНЕНИЕ ЦЕН, СКРАПИНГ, ВЕБ-ПРИЛОЖЕНИЕ, FASTAPI, NEXT.JS, GEMINI, OPENROUTER, "
    "SSE, ПОТОКОВАЯ ПЕРЕДАЧА ДАННЫХ, ИСТОРИЯ ЦЕН, АГРЕГАТОР ЦЕН, ЭЛЕКТРОННАЯ "
    "КОММЕРЦИЯ, DOCKER, POSTGRESQL, АНТИБОТ-ЗАЩИТА.",
    "Объект исследования — процессы поиска и сравнения цен на товары на "
    "маркетплейсах России и Беларуси; предмет исследования — методы "
    "автоматизированного метапоиска и интеллектуального анализа цен. Цель работы — "
    "проектирование и разработка веб-платформы метапоиска с интеграцией "
    "AI-ассистента на базе модели Gemini 2.5 Flash и потоковой передачей "
    "результатов через Server-Sent Events.",
    "В первой главе проанализированы рынок электронной коммерции России и Беларуси "
    "и существующие решения, выявлены функциональные пробелы. Во второй главе "
    "спроектирована архитектура системы: модель базы данных, REST API, "
    "SSE-стриминг и контейнеризация в Docker. В третьей главе реализованы парсеры "
    "маркетплейсов, AI-коррекция поисковых запросов и серверная часть на FastAPI. "
    "В четвёртой главе проведены функциональное и нагрузочное тестирование и "
    "оптимизация производительности.",
    "В результате разработана и развёрнута в промышленную эксплуатацию система "
    "Smart Price (smrt-price.ru), обеспечивающая мгновенный поиск, сравнение цен "
    "на маркетплейсах и AI-рекомендации по выбору товаров.",
]


def process(fn):
    d = Document(fn)
    paras = d.paragraphs
    ref = None
    end = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == "РЕФЕРАТ":
            ref = i
        elif ref is not None and t in ("СОДЕРЖАНИЕ", "ОГЛАВЛЕНИЕ"):
            end = i
            break
    # реферат body = non-empty paragraphs between ref+1 and end
    body = [p for p in paras[ref + 1 : end] if p.text.strip()]
    # write NEW into first len(NEW); delete the rest
    for k, txt in enumerate(NEW):
        p = body[k]
        if p.runs:
            p.runs[0].text = txt
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(txt)
    for p in body[len(NEW) :]:
        p._element.getparent().remove(p._element)
    d.save(fn)
    return {"ref": ref, "end": end, "body_was": len(body), "now": len(NEW)}


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "->", process(fn))
