# -*- coding: utf-8 -*-
"""Script 1 — fonts & sizes & table header.
   - Every run (body + tables + title) -> Times New Roman (kill Consolas/theme).
   - Body runs from РЕФЕРАТ onward (NOT title page, NOT table cells) -> 14 pt.
   - Heading styles 1 & 2 -> 14 pt.
   - Table header row (row 0) -> bold.
   - Table body keeps its smaller size (allowed); only font family normalized.
   - Table captions -> distributed across full width (handled here: alignment).
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TNR = "Times New Roman"


def force_tnr(run):
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for a in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        if rf.get(qn("w:" + a)) is not None:
            del rf.attrib[qn("w:" + a)]
    for a in ("ascii", "hAnsi", "cs", "eastAsia"):
        rf.set(qn("w:" + a), TNR)


def ptext(e):
    return "".join(t.text or "" for t in e.findall(".//" + qn("w:t"))).strip()


def process(fn):
    d = Document(fn)
    # index of РЕФЕРАТ (title page = everything before it)
    ref = None
    for i, p in enumerate(d.paragraphs):
        if p.text.strip() == "РЕФЕРАТ":
            ref = i
            break

    # 1. global TNR on body paragraphs
    for p in d.paragraphs:
        for r in p.runs:
            force_tnr(r)
    # 1b. global TNR inside tables
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        force_tnr(r)

    # 2. body runs from РЕФЕРАТ onward -> 14 pt (skip title page + table cells)
    if ref is not None:
        for i, p in enumerate(d.paragraphs):
            if i < ref:
                continue
            for r in p.runs:
                r.font.size = Pt(14)

    # 3. heading styles 1 & 2 -> 14 pt
    for nm in ("Heading 1", "Heading 2"):
        rpr = d.styles[nm].element.get_or_add_rPr()
        for tag in ("w:sz", "w:szCs"):
            e = rpr.find(qn(tag))
            if e is None:
                e = OxmlElement(tag)
                rpr.append(e)
            e.set(qn("w:val"), "28")

    # 4. table header row bold + table captions distributed
    for t in d.tables:
        if t.rows:
            for cell in t.rows[0].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    # captions: alignment distribute (size already 14 from body pass)
    body = d.element.body
    ch = list(body.iterchildren())
    from docx.text.paragraph import Paragraph

    caps = 0
    for i, e in enumerate(ch):
        if e.tag != qn("w:tbl"):
            continue
        j = i - 1
        while j >= 0:
            if ch[j].tag == qn("w:p") and ptext(ch[j]):
                if ptext(ch[j]).startswith("Таблица"):
                    pp = Paragraph(ch[j], d)
                    pp.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
                    caps += 1
                break
            j -= 1

    d.save(fn)
    return {"referat_idx": ref, "captions_distributed": caps}


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "->", process(fn))
