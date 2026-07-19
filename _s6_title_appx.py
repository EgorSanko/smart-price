# -*- coding: utf-8 -*-
"""Script 6 — uniform title signature underlines + appendix headings out of TOC."""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn

NAMES = ("Санько", "Вагнер", "Чалова")


def process(fn):
    d = Document(fn)

    # --- title: make signature underlines uniform (extend to 16.5 cm) ---
    sig = 0
    for p in d.paragraphs[:30]:
        t = p.text.strip()
        if any(n in t for n in NAMES):
            # remove existing tab stops, add a right tab w/ underscore leader at right margin
            pf = p.paragraph_format
            pPr = p._element.get_or_add_pPr()
            old = pPr.find(qn("w:tabs"))
            if old is not None:
                pPr.remove(old)
            pf.tab_stops.add_tab_stop(
                Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES
            )
            # append a trailing tab so the leader draws the line to 16.5 cm
            if not p.text.endswith("\t"):
                p.add_run("\t")
            sig += 1

    # --- appendix А/Б -> Normal (keep look) so only ПРИЛОЖЕНИЯ stays in TOC ---
    appx = 0
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith("ПРИЛОЖЕНИЕ А") or t.startswith("ПРИЛОЖЕНИЕ Б"):
            brk = not t.startswith("ПРИЛОЖЕНИЕ А")  # А shares divider page, Б breaks
            p.style = d.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(14)
            p.paragraph_format.page_break_before = brk
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            # drop any outlineLvl leftover
            pPr = p._element.get_or_add_pPr()
            for ol in pPr.findall(qn("w:outlineLvl")):
                pPr.remove(ol)
            appx += 1

    d.save(fn)
    return {"signatures": sig, "appendix_to_normal": appx}


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "->", process(fn))
