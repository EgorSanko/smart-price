# -*- coding: utf-8 -*-
"""Fix-up patch: correct wrongly-placed conclusion sentence, kill last Яндекс.Маркет."""
import sys
from docx import Document

SRC = r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx"
doc = Document(SRC)

CHEAPER = " Реализована функция «Найти дешевле», позволяющая по ссылке на товар Яндекс Маркета подобрать более дешёвые предложения у сторонних продавцов через асинхронную Celery-задачу с извлечением данных из вложенной структуры EAliceOfferList и детектированием отказов по регулярному выражению."


def replace_all_runs(paragraph, old, new):
    """Aggressive replace: works even when old string is split across runs."""
    if old not in paragraph.text:
        return False
    # Try per-run first
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
    if old not in paragraph.text:
        return True
    # Collapse: put full replaced text into the first run
    full = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


# --- 1. Strip mis-appended sentence from TOC line (para 62) ---
for p in doc.paragraphs:
    if CHEAPER.strip() in p.text and p.text.lstrip().startswith("4 ТЕСТИРОВАНИЕ"):
        replace_all_runs(p, CHEAPER, "")
        print("Stripped sentence from TOC line")
        break

# --- 2. Append sentence to the real conclusion body paragraph ---
# it's the paragraph right after body heading '3.14 Выводы по третьей главе'
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "3.14 Выводы по третьей главе":
        target = doc.paragraphs[i + 1]
        # avoid double-append
        if CHEAPER.strip() not in target.text:
            if target.runs:
                target.runs[-1].text = target.runs[-1].text.rstrip() + CHEAPER
            else:
                target.add_run(CHEAPER)
            print(f"Appended cheaper mention to body para {i+1}")
        break

# --- 3. Final sweep for Яндекс.Маркет ---
cnt = 0
for p in doc.paragraphs:
    if "Яндекс.Маркет" in p.text:
        replace_all_runs(p, "Яндекс.Маркет", "Яндекс Маркет")
        cnt += 1
print(f"Stale Яндекс.Маркет fixed: {cnt}")

doc.save(SRC)
print("Saved.")
