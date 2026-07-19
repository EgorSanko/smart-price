# -*- coding: utf-8 -*-
"""Script 3 — italicize every Latin (anglicism) token across the whole document,
   INCLUDING underscores/dots/hyphens inside identifiers (user_id, Next.js,
   Server-Sent, FastAPI, AI...). Cyrillic stays upright. Run-splitting.
"""
import copy
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

XSPACE = "{http://www.w3.org/XML/1998/namespace}space"
# a Latin token: starts with a Latin letter, may contain digits and ._- joiners
LAT = re.compile(r"[A-Za-z][A-Za-z0-9]*(?:[._\-/][A-Za-z0-9]+)*")


def italicize_run(run_el):
    # skip runs that carry a drawing/image
    if run_el.find(".//" + qn("w:drawing")) is not None:
        return 0
    t_el = run_el.find(qn("w:t"))
    if t_el is None or not t_el.text:
        return 0
    text = t_el.text
    spans = [(m.start(), m.end()) for m in LAT.finditer(text)]
    if not spans:
        return 0
    # build segments
    segs = []
    idx = 0
    for s, e in spans:
        if s > idx:
            segs.append((text[idx:s], False))
        segs.append((text[s:e], True))
        idx = e
    if idx < len(text):
        segs.append((text[idx:], False))
    # whole run is one latin token already italic -> leave
    rPr = run_el.find(qn("w:rPr"))
    already_it = (
        rPr is not None
        and rPr.find(qn("w:i")) is not None
        and rPr.find(qn("w:i")).get(qn("w:val")) not in ("0", "false")
    )
    if len(segs) == 1 and segs[0][1] and already_it:
        return 0
    new_runs = []
    for seg_text, is_lat in segs:
        nr = copy.deepcopy(run_el)
        for old_t in nr.findall(qn("w:t")):
            nr.remove(old_t)
        nt = OxmlElement("w:t")
        nt.set(XSPACE, "preserve")
        nt.text = seg_text
        nr.append(nt)
        npr = nr.find(qn("w:rPr"))
        if npr is None:
            npr = OxmlElement("w:rPr")
            nr.insert(0, npr)
        i_el = npr.find(qn("w:i"))
        ics = npr.find(qn("w:iCs"))
        if is_lat:
            if i_el is None:
                i_el = OxmlElement("w:i")
                npr.append(i_el)
            i_el.set(qn("w:val"), "true")
            if ics is None:
                ics = OxmlElement("w:iCs")
                npr.append(ics)
            ics.set(qn("w:val"), "true")
        else:
            if i_el is not None:
                npr.remove(i_el)
            if ics is not None:
                npr.remove(ics)
        new_runs.append(nr)
    for nr in new_runs:
        run_el.addprevious(nr)
    run_el.getparent().remove(run_el)
    return 1


def process(fn):
    d = Document(fn)
    n = 0
    # body paragraphs
    for p in d.paragraphs:
        for r_el in list(p._element.findall(qn("w:r"))):
            n += italicize_run(r_el)
    # table cells
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r_el in list(p._element.findall(qn("w:r"))):
                        n += italicize_run(r_el)
    d.save(fn)
    return n


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "runs italicized:", process(fn))
