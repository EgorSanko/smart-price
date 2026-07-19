# -*- coding: utf-8 -*-
"""Dump every occurrence of Russia/Belarus with surrounding context."""
import re
import io
from docx import Document

doc = Document(r"C:/Users/egor3/Desktop/smart-price/ДИПЛОМ.docx")


def iter_all_paragraphs(d):
    for i, p in enumerate(d.paragraphs):
        yield ("P", i, p.text, p)
    for ti, t in enumerate(d.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield (f"t{ti}r{ri}c{ci}p{pi}", None, p.text, p)


out = io.StringIO()

for kind, idx, text, p in iter_all_paragraphs(doc):
    if not text:
        continue
    for m in re.finditer(r".{0,40}(Росси|Белару|Беларус|Россий).{0,40}", text):
        out.write(f"[{kind}{idx}] {m.group(0)!r}\n")

# Also dump per-run text where Russia/Belarus appears, to see if runs are split weirdly
out.write("\n\n=== Per-run dump for paragraphs with Russia/Belarus ===\n")
for kind, idx, text, p in iter_all_paragraphs(doc):
    if not text:
        continue
    if "Росси" in text or "Беларус" in text or "Белару" in text:
        out.write(f"\n[{kind}{idx}]\n")
        for ri, r in enumerate(p.runs):
            out.write(f"  run{ri}: {r.text!r}\n")

with open(
    r"C:/Users/egor3/Desktop/smart-price/_russia_belarus.txt", "w", encoding="utf-8"
) as f:
    f.write(out.getvalue())
print("Done")
