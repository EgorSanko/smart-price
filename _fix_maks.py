# -*- coding: utf-8 -*-
"""Maks' 3 fixes:
  1. Subsection headings (Heading 2, x.x) -> NOT bold (chapters stay bold).
  2. Empty paragraph (spacing) after each table.
  3. Tables -> left alignment so the table left edge matches the «Таблица N» caption.
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def process(fn):
    d = Document(fn)
    s = {"h2": 0, "after_tbl": 0, "tbl_left": 0}

    # 1. Heading 2 not bold (style + clear direct bold on its runs)
    try:
        st = d.styles["Heading 2"]
        st.font.bold = False
        rpr = st.element.get_or_add_rPr()
        for tag in ("w:b", "w:bCs"):
            e = rpr.find(qn(tag))
            if e is not None:
                rpr.remove(e)
        # explicit off so it never inherits bold
        b = OxmlElement("w:b")
        b.set(qn("w:val"), "0")
        rpr.append(b)
        bcs = OxmlElement("w:bCs")
        bcs.set(qn("w:val"), "0")
        rpr.append(bcs)
        s["h2"] = 1
    except KeyError:
        pass
    # clear any direct bold on H2 paragraph runs
    for p in d.paragraphs:
        if p.style.name == "Heading 2":
            for r in p.runs:
                rPr = r._element.find(qn("w:rPr"))
                if rPr is not None:
                    for tag in ("w:b", "w:bCs"):
                        e = rPr.find(qn(tag))
                        if e is not None:
                            rPr.remove(e)

    # 3. tables left-aligned + tblInd 0
    for t in d.tables:
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        tblPr = t._tbl.tblPr
        ind = tblPr.find(qn("w:tblInd"))
        if ind is not None:
            ind.set(qn("w:w"), "0")
            ind.set(qn("w:type"), "dxa")
        s["tbl_left"] += 1

    # 2. empty paragraph after each table (if not already present)
    for t in d.tables:
        tbl = t._tbl
        nxt = tbl.getnext()
        need = True
        if nxt is not None and nxt.tag == qn("w:p"):
            txt = "".join(x.text or "" for x in nxt.findall(".//" + qn("w:t"))).strip()
            # also treat a paragraph that has a drawing as content
            if txt == "" and not nxt.findall(".//" + qn("w:drawing")):
                need = False
        if need:
            p = OxmlElement("w:p")
            tbl.addnext(p)
            s["after_tbl"] += 1

    d.save(fn)
    return s


for fn in ["ДИПЛОМ_нормоконтроль.docx", "ВКР_Санько_итог.docx"]:
    print(fn, "->", process(fn))
